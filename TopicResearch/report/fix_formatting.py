"""
专业排版修复脚本 (安全版 v2)
- 数据修正: 基于 ifind/FRED 验证结果修正文档中的错误数据
- 格式修复: 14 项专业排版
- safe_save: 先写临时文件，再 shutil.move 覆盖
"""
import os
import re
import shutil
import tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ============================================================
# 数据修正规则
# ============================================================
# 格式: (旧文本, 新文本, 描述, 数据来源)
DATA_FIXES = [
    # 中国 M2 2007: 18.5% → 16.7% (ifind EDB: 2007-12=16.7%)
    ('18.5%', '16.7%', 'China M2 2007', 'ifind EDB M001625222'),

    # 美国 M2 2006: 4.9% → 5.9% (FRED M2SL: Dec-on-Dec)
    ('4.9%', '5.9%', 'US M2 2006', 'FRED M2SL'),

    # NVDA: 800% → 1,035% (ifind: $16.92→$192)
    ('800%', '1,035%', 'NVDA return', 'ifind global_stock'),

    # NVDA: $130 → $192 (ifind: current price)
    ('$130', '$192', 'NVDA current price', 'ifind global_stock'),

    # 格林斯潘: +0.65% → 接近零 (FRED: FF-GS10=-0.25%)
    ('+0.65%', '接近零', 'Greenspan spread endpoint', 'FRED GS10+FEDFUNDS'),
]


def safe_save(doc, target_path):
    """安全保存: 先写临时文件，再原子替换"""
    dir_name = os.path.dirname(os.path.abspath(target_path))
    fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=dir_name)
    os.close(fd)
    try:
        doc.save(tmp_path)
        shutil.move(tmp_path, target_path)
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        tcPr.remove(shd)
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tcPr.append(shd)


def set_table_borders(table, color='CBD5E1'):
    """设置表格统一边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_full_width(table):
    """表格宽度 100%"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')


def fix_data_in_paragraphs(paragraphs, log):
    """在段落 runs 中修正数据"""
    count = 0
    for p in paragraphs:
        for r in p.runs:
            original = r.text
            modified = original
            for old, new, desc, source in DATA_FIXES:
                if old in modified:
                    modified = modified.replace(old, new)
                    log.append(f'  {desc}: "{old}"→"{new}" ({source})')
            if modified != original:
                r.text = modified
                count += 1
    return count


def fix_data(doc):
    """修正文档中的错误数据"""
    log = []

    # 修正段落
    p_count = fix_data_in_paragraphs(doc.paragraphs, log)

    # 修正表格
    t_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t_count += fix_data_in_paragraphs(cell.paragraphs, log)

    print(f'[数据修正] 段落: {p_count}, 表格: {t_count}, 总修改: {len(log)}')
    for entry in log:
        print(entry)

    return log


def apply_formatting(doc):
    """应用 14 项格式修复"""

    # === 1. 页边距 ===
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    print('[1/14] 页边距')

    # === 2. Normal 样式 ===
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.line_spacing = 1.7
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Pt(22)
    style.paragraph_format.space_before = Pt(0)
    print('[2/14] Normal 样式')

    # === 3. Normal (Web) -> Normal ===
    web_count = 0
    for p in doc.paragraphs:
        if p.style.name == 'Normal (Web)':
            p.style = doc.styles['Normal']
            web_count += 1
    print(f'[3/14] Normal (Web)->Normal: {web_count}')

    # === 4-7. 标题样式 ===
    cfg = {
        1: (18, '1E3A5F', 24, 16),
        2: (14, '1E3A5F', 18, 12),
        3: (12, '1E3A5F', 12, 8),
        4: (11, '475569', 8, 6),
    }
    for lv, (sz, clr, sb, sa) in cfg.items():
        s = doc.styles[f'Heading {lv}']
        s.font.name = 'Times New Roman'
        s.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        s.font.size = Pt(sz)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(clr)
        s.paragraph_format.first_line_indent = Pt(0)
        s.paragraph_format.line_spacing = 1.3
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after = Pt(sa)
    print('[4-7/14] 标题 H1-H4')

    # === 8-9. 正文行距+缩进 ===
    for p in doc.paragraphs:
        if p.style.name == 'Normal':
            p.paragraph_format.line_spacing = 1.7
            p.paragraph_format.first_line_indent = Pt(22)
            p.paragraph_format.space_after = Pt(6)
    print('[8-9/14] 正文行距+缩进')

    # === 10. 段后间距 ===
    print('[10/14] 段后间距')

    # === 11-13. 表格 ===
    for table in doc.tables:
        set_table_borders(table)
        set_table_full_width(table)
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.line_spacing = 1.3

                if ri == 0:
                    set_cell_shading(cell, '1E3A5F')
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                            r.font.size = Pt(10)
                            r.font.name = 'Times New Roman'
                            r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                elif ri % 2 == 0:
                    set_cell_shading(cell, 'F1F5F9')
                else:
                    set_cell_shading(cell, 'FFFFFF')

                for p in cell.paragraphs:
                    for r in p.runs:
                        if not r.font.size:
                            r.font.size = Pt(9.5)
                        r.font.name = 'Times New Roman'
                        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    print(f'[11-14/14] 表格 ({len(doc.tables)})')

    return doc


def verify(doc, fix_log):
    """验证修复结果"""
    print('\n=== 验证 ===')
    print(f'段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}')

    # 样式检查
    web = sum(1 for p in doc.paragraphs if p.style.name == 'Normal (Web)')
    print(f'Normal(Web)残留: {web}')

    # 数据修正验证
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text

    checks = {
        'M2 18.5% (should be 16.7%)': '18.5%' in all_text,
        'NVDA 800% (should be 1035%)': '800%' in all_text,
        'NVDA $130 (should be $192)': '$130' in all_text,
        'M2 4.9% (should be 5.9%)': '4.9%' in all_text,
        '格潘 +0.65% (should be 接近零)': '+0.65%' in all_text,
    }

    all_ok = True
    for name, still_exists in checks.items():
        if still_exists:
            print(f'  ⚠️ 未修正: {name}')
            all_ok = False
        else:
            print(f'  ✅ 已修正: {name}')

    return all_ok


if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1:
        path = sys.argv[1]
    else:
        path = '/home/ll/llmwikify/TopicResearch/report/复盘05-07有色vs当前AI板块投资研究.docx'

    print(f'加载: {path}')
    doc = Document(path)
    print(f'  段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}')

    # Step 1: 数据修正
    print('\n=== 数据修正 ===')
    fix_log = fix_data(doc)

    # Step 2: 格式修复
    print('\n=== 格式修复 ===')
    apply_formatting(doc)

    # Step 3: 验证
    ok = verify(doc, fix_log)

    # Step 4: 保存
    safe_save(doc, path)
    print(f'\n✅ 已保存: {path}')
    print(f'   大小: {os.path.getsize(path)/1024:.1f} KB')
    print(f'   验证: {"通过" if ok else "有未修正项"}')
