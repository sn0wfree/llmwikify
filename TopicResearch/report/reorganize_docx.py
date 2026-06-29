"""
重新组织文档：将AI IPO抽血分析移到1.3.4章节后面
"""
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCX_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'
OUTPUT_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'
CHART_DIR = '/home/ll/llmwikify/TopicResearch/report/charts'

def main():
    print('Loading document...')
    doc = docx.Document(DOCX_PATH)
    
    # 识别所有段落及其类型
    paragraphs_info = []
    for i, p in enumerate(doc.paragraphs):
        info = {
            'index': i,
            'text': p.text,
            'style': p.style.name if p.style else 'Normal',
            'has_image': False,
        }
        # 检查是否有图片
        for run in p.runs:
            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                info['has_image'] = True
                break
        paragraphs_info.append(info)
    
    # 找到关键位置
    s134_start = None  # 1.3.4 开始
    s134_end = None    # 1.3.4 结束（下一个Heading）
    ipo_supplement_start = None  # 补充：本轮AI巨型IPO抽血分析
    ipo_supplement_end = None
    s2_supplement_start = None  # 补充：S2信号详细分析
    appendix_d_start = None  # 附录 D
    
    for i, info in enumerate(paragraphs_info):
        if '1.3.4' in info['text'] and info['style'].startswith('Heading'):
            s134_start = i
        if s134_start and i > s134_start and info['style'].startswith('Heading') and '1.4' in info['text']:
            s134_end = i
            break
        if '补充：本轮AI巨型IPO抽血分析' in info['text']:
            ipo_supplement_start = i
        if ipo_supplement_start and '补充：S2信号详细分析' in info['text']:
            ipo_supplement_end = i
            break
    
    print(f'1.3.4 section: {s134_start} to {s134_end}')
    print(f'IPO supplement: {ipo_supplement_start} to {ipo_supplement_end}')
    
    # 创建新文档
    new_doc = docx.Document()
    
    # 复制段落到新文档，调整顺序
    # 1. 复制 1.3.4 之前的内容
    print('\nCopying content before 1.3.4...')
    for i in range(s134_start):
        p = paragraphs_info[i]
        new_p = new_doc.add_paragraph()
        if p['style'] != 'Normal':
            try:
                new_p.style = p['style']
            except:
                pass
        run = new_p.add_run(p['text'])
        run.font.size = Pt(9)
    
    # 2. 复制 1.3.4 内容
    print('Copying 1.3.4 section...')
    for i in range(s134_start, s134_end):
        p = paragraphs_info[i]
        new_p = new_doc.add_paragraph()
        if p['style'] != 'Normal':
            try:
                new_p.style = p['style']
            except:
                pass
        run = new_p.add_run(p['text'])
        run.font.size = Pt(9)
    
    # 3. 插入 AI IPO 抽血分析（在 1.3.4 后面）
    print('Inserting AI IPO analysis after 1.3.4...')
    
    # 添加标题
    p = new_doc.add_paragraph()
    run = p.add_run('1.3.4-AI IPO抽血分析（本轮AI巨型IPO）')
    run.font.size = Pt(10)
    run.font.bold = True
    
    # 添加背景说明
    p = new_doc.add_paragraph()
    run = p.add_run(
        '2026年6月，SpaceX在纳斯达克成功上市（股票代码SPCX），募资750亿美元，创下全球资本市场史上最大IPO纪录。'
        '紧随其后，OpenAI计划于2026Q4-2027Q1进行IPO，规模预计1000亿美元。'
        '本轮AI巨兽集中上市的"抽血效应"值得高度关注。'
    )
    run.font.size = Pt(9)
    
    # 添加表格1：本轮AI IPO全景
    p = new_doc.add_paragraph()
    run = p.add_run('表 1.3-1：本轮AI巨型IPO全景')
    run.font.size = Pt(9)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = new_doc.add_table(rows=5, cols=5)
    headers = ['公司', '状态', '募资规模', '估值', '抽血占比']
    rows = [
        ['SpaceX', '✅ 已上市 (2026.6.12)', '$750亿', '$1.77万亿', '0.17%'],
        ['OpenAI', '拟IPO (2026Q4-2027Q1)', '$1000亿', '$8500亿', '0.22%'],
        ['Anthropic', '拟IPO (2027+)', '$500亿', '$2000亿', '0.11%'],
        ['合计', '-', '$2250亿', '-', '0.50%'],
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # 添加表格2：SpaceX vs 中石油
    p = new_doc.add_paragraph()
    run = p.add_run('表 1.3-2：SpaceX vs 中石油 IPO对比')
    run.font.size = Pt(9)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table2 = new_doc.add_table(rows=8, cols=4)
    headers2 = ['维度', '2007 中石油', '2026 SpaceX', '差异']
    rows2 = [
        ['募资规模', '¥668亿 ($97亿)', '$750亿', 'SpaceX是中石油7.7倍'],
        ['上市估值', '$1.08万亿', '$1.77万亿', 'SpaceX高64%'],
        ['首日涨幅', '+191% (开盘)', '+29% (开盘)', '中石油更疯狂'],
        ['冻结资金', '¥3.3万亿 (40%)', '$2500亿认购 (4倍)', '中石油抽血更猛'],
        ['散户参与', '>60%', '30%', '散户占比下降'],
        ['盈利状态', '盈利', '亏损 ($49.4亿)', 'SpaceX尚未盈利'],
        ['后续表现', '一年内破发', '待观察', '-'],
    ]
    for j, h in enumerate(headers2):
        cell = table2.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(8)
    for r, row_data in enumerate(rows2):
        for c, val in enumerate(row_data):
            cell = table2.rows[r+1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # 添加抽血机制分析
    p = new_doc.add_paragraph()
    run = p.add_run('AI IPO抽血机制分析：')
    run.font.size = Pt(9)
    run.font.bold = True
    
    analyses = [
        ('（1）抽血规模对比：', 'SpaceX募资750亿美元，占标普500流通市值约0.17%；OpenAI预计募资1000亿美元，占纳斯达克100流通市值约0.4%。合计潜在抽血2250亿美元，占标普500流通市值约0.50%。对比2007年中石油冻结3.3万亿元（占A股流通市值40%），本轮抽血规模远小于2007年。'),
        ('（2）存量占用更严重：', '本轮AI牛市中，NVIDIA+Microsoft已占纳斯达克100约24%权重。SpaceX上市后迅速纳入纳斯达克100，进一步加剧头部集中。当AI巨兽占据指数主要权重时，新IPO的边际抽血效应可能被放大。'),
        ('（3）流动性环境差异：', '2007年M2增速17%+外汇占款被动投放，流动性极度宽松；2026年M2增速8%+美联储维持高利率，流动性相对紧张。这意味着本轮IPO对市场流动性的冲击可能更敏感。'),
        ('（4）散户参与度下降：', 'SpaceX散户份额30%（远高于近年平均水平），但中石油时代散户占比超60%。机构主导的IPO通常波动更小，但一旦下跌，机构止损盘可能更猛烈。'),
    ]
    
    for title, content in analyses:
        p = new_doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(9)
        run.font.bold = True
        run = p.add_run(content)
        run.font.size = Pt(9)
    
    # 添加关键判断
    p = new_doc.add_paragraph()
    run = p.add_run('关键判断：')
    run.font.size = Pt(9)
    run.font.bold = True
    
    p = new_doc.add_paragraph()
    run = p.add_run(
        '本轮AI IPO抽血规模虽小于2007年，但需高度关注两个风险点：'
        '（1）多只AI巨兽集中上市的时间窗口重叠（SpaceX已上市+OpenAI即将IPO）；'
        '（2）存量占用严重（NVIDIA+Microsoft+SpaceX已占纳斯达克100约26%）。'
        '核心监测指标：OpenAI IPO后30日表现（S2信号触发点）。'
    )
    run.font.size = Pt(9)
    
    # 4. 复制 1.3.4 之后到文档末尾的内容（跳过补充章节）
    print('Copying remaining content (skipping supplements)...')
    
    # 找到第一个补充章节的位置
    first_supplement = len(paragraphs_info)
    for i in range(s134_end, len(paragraphs_info)):
        if '补充：' in paragraphs_info[i]['text'] or '附录 D' in paragraphs_info[i]['text']:
            first_supplement = i
            break
    
    print(f'First supplement at index: {first_supplement}')
    
    for i in range(s134_end, first_supplement):
        p = paragraphs_info[i]
        new_p = new_doc.add_paragraph()
        if p['style'] != 'Normal':
            try:
                new_p.style = p['style']
            except:
                pass
        run = new_p.add_run(p['text'])
        run.font.size = Pt(9)
    
    # 6. 添加附录 D（图表）
    print('Adding appendix D with charts...')
    new_doc.add_page_break()
    p = new_doc.add_heading('附录 D：宏观经济背景图表', level=1)
    
    charts = [
        ('chart_gdp_nominal.png', '图 1.1：名义 GDP 增速对比（2000-2026）', '2005-2007 年中国名义 GDP 增速从 15.6% 升至 23.2%。数据来源：ifind EDB。'),
        ('chart_gdp_real.png', '图 1.2：实际 GDP 增速对比（2000-2026）', '剔除价格因素后，中国实际 GDP 增速 11-14%。数据来源：ifind EDB。'),
        ('chart_cpi.png', '图 1.3：CPI 通胀对比（2000-2026）', '2005-2007 年中国 CPI 从 1.8% 升至 4.8%。数据来源：ifind EDB。'),
        ('chart_monetary.png', '图 1.4：中国货币政策背离', '2006-2007 年 PBOC 加息 6 次+提 RRR 10 次，但 M2 维持 17-19%。数据来源：ifind EDB。'),
        ('chart_term_spread.png', '图 1.5：市场利率与货币政策类型', '美国价格型、中国数量型→价格型、日本 YCC。数据来源：ifind EDB。'),
        ('chart_fx_trade.png', '图 1.6：汇率与贸易', '2005.7.21 汇改，DXY 从 120 跌至 80。数据来源：ifind EDB。'),
        ('chart_forex_reserve.png', '图 1.7：外汇储备/GDP', '中国外储/GDP 从 15% 升至 49%（2014）。数据来源：ifind EDB。'),
        ('chart_impossible_trinity.png', '图 1.8：蒙代尔不可能三角', '固定汇率、资本自由流动、独立货币政策三者不可兼得。'),
    ]
    
    for filename, title, description in charts:
        chart_path = os.path.join(CHART_DIR, filename)
        if not os.path.exists(chart_path):
            print(f'  ⚠️ Chart not found: {chart_path}')
            continue
        
        p = new_doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(10)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = new_doc.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(chart_path, width=Inches(6))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f'  ✅ Added: {title}')
        except Exception as e:
            print(f'  ❌ Failed: {e}')
        
        p = new_doc.add_paragraph()
        run = p.add_run(description)
        run.font.size = Pt(9)
    
    # 保存
    print(f'\nSaving to {OUTPUT_PATH}...')
    new_doc.save(OUTPUT_PATH)
    print(f'✅ Done! Saved to {OUTPUT_PATH}')
    
    # 验证
    final_doc = docx.Document(OUTPUT_PATH)
    rels = final_doc.part.rels
    image_count = sum(1 for rel in rels.values() if 'image' in rel.reltype)
    print(f'\nVerification: {len(final_doc.paragraphs)} paragraphs, {image_count} images')

if __name__ == '__main__':
    main()
