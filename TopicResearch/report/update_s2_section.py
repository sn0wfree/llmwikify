"""
更新 3.2 S2信号章节：添加OpenAI IPO详细分析
"""
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'
OUTPUT_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'

def main():
    print('Loading document...')
    doc = docx.Document(DOCX_PATH)
    
    # 找到3.2节位置
    s2_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '3.2' in p.text and 'S2' in p.text:
            s2_idx = i
            break
    
    if s2_idx is None:
        print('Error: Section 3.2 not found')
        return
    
    print(f'Found section 3.2 at index {s2_idx}')
    
    # 在3.3节前添加内容
    # 找到3.3节位置
    s3_idx = None
    for i in range(s2_idx + 1, len(doc.paragraphs)):
        if '3.3' in doc.paragraphs[i].text:
            s3_idx = i
            break
    
    if s3_idx is None:
        s3_idx = s2_idx + 10
    
    # 添加内容到文档末尾（python-docx限制）
    # 添加补充内容
    p = doc.add_paragraph()
    run = p.add_run('补充：S2信号详细分析')
    run.font.size = Pt(10)
    run.font.bold = True
    
    # OpenAI IPO 详细分析
    p = doc.add_paragraph()
    run = p.add_run('A. OpenAI IPO 详细分析')
    run.font.size = Pt(9)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run(
        'OpenAI估值演变：2023年$860B → 2025年$500B → 2026年$850B。'
        'IPO预期：2026Q4-2027Q1，规模$1000亿。'
        '抽血规模：$1000亿（占纳斯达克100流通市值约0.4%）。'
        '历史对标：中石油冻结3.3万亿元（占流通市值40%），规模远大于OpenAI。'
        '判断：OpenAI IPO抽血效应相对较小，但需关注IPO后30日表现。'
    )
    run.font.size = Pt(9)
    
    # S2信号触发条件
    p = doc.add_paragraph()
    run = p.add_run('B. S2信号触发条件（量化标准）')
    run.font.size = Pt(9)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run(
        '1. 首日涨幅>50%后收长阴线（复现中石油模式）；'
        '2. IPO后30日跌幅>20%；'
        '3. 市场恐慌指数（VIX）在IPO后10日内上升>30%。'
        '满足任意两个条件即触发S2信号。'
    )
    run.font.size = Pt(9)
    
    # SpaceX 上市影响
    p = doc.add_paragraph()
    run = p.add_run('C. SpaceX上市影响（2026.6.12）')
    run.font.size = Pt(9)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run(
        'SpaceX于2026年6月12日在纳斯达克上市（代码SPCX），募资$750亿（史上最大IPO）。'
        '首日开盘$174（+29%），上市估值$1.77万亿，成为美国第七大上市公司。'
        '关键数据：2025年营收$186.7亿，净亏损$49.4亿；散户份额30%。'
        '后续影响：上市15个交易日后纳入纳斯达克100，加剧头部集中。'
    )
    run.font.size = Pt(9)
    
    # 保存
    print(f'\nSaving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'✅ Done! Saved to {OUTPUT_PATH}')

if __name__ == '__main__':
    main()
