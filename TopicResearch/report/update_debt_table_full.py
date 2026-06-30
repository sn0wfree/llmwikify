"""
更新文档表格9（三国债务结构）完整21列数据（2005-2025）
"""
import os
from lxml import etree
from docx import Document

BASE_DIR = os.path.expanduser("~/llmwikify/TopicResearch/report")
V210_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.10.docx")
OUT_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.11.docx")

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# 完整数据（2005-2025）
YEARS = list(range(2005, 2026))

# 数据
US_GOV_DEBT = [61.6, 62.0, 62.6, 67.5, 82.1, 90.9, 95.5, 99.6, 100.2, 102.1, 99.8, 104.8, 104.3, 105.0, 106.5, 126.1, 123.2, 120.6, 120.4, 121.7, 123.3]
US_HH_LEVERAGE = [92.7, 96.9, 98.7, 96.1, 97.0, 91.7, 87.8, 83.6, 81.8, 79.0, 77.3, 77.4, 77.0, 75.0, 74.6, 77.9, 76.7, 74.5, 71.8, 69.3, 68.1]
US_CORP_LEVERAGE = [64.8, 67.2, 72.1, 74.6, 72.6, 69.0, 68.3, 69.1, 69.9, 71.7, 73.7, 75.2, 77.3, 78.2, 78.3, 86.3, 82.8, 78.9, 75.3, 73.3, 72.2]

CN_GOV_DEBT = [25.9, 25.2, 28.7, 26.7, 34.0, 33.3, 33.2, 33.8, 36.4, 39.3, 40.8, 49.7, 53.9, 55.6, 59.4, 69.0, 70.1, 77.3, 84.1, 90.4, 99.2]
CN_HH_LEVERAGE = [16.6, 17.2, 18.5, 17.6, 23.1, 26.9, 27.5, 29.5, 32.9, 35.3, 38.5, 43.8, 47.8, 51.2, 55.0, 61.1, 60.6, 60.7, 61.9, 61.4, 59.4]
CN_CORP_LEVERAGE = [98.4, 97.6, 94.6, 93.7, 113.6, 118.4, 116.1, 126.0, 133.5, 139.8, 148.5, 154.4, 153.8, 147.9, 149.0, 159.2, 150.9, 157.7, 163.9, 168.4, 174.6]

JP_GOV_DEBT = [174.6, 174.1, 173.0, 180.9, 198.8, 205.9, 219.2, 226.1, 229.5, 233.3, 228.3, 232.4, 231.3, 232.4, 236.4, 258.4, 253.7, 227.8, 220.3, 214.5, 206.5]
JP_HH_LEVERAGE = [62.1, 61.1, 60.0, 60.2, 63.5, 61.2, 61.8, 61.4, 61.3, 60.7, 59.5, 59.8, 60.2, 61.2, 62.3, 67.2, 64.8, 64.7, 62.4, 62.1, 61.1]
JP_CORP_LEVERAGE = [310.9, 307.3, 303.9, 313.9, 339.3, 337.7, 351.0, 356.5, 360.6, 367.2, 360.0, 368.1, 368.2, 373.0, 382.9, 422.5, 415.2, 396.0, 384.3, 371.3, 353.9]


def get_element_text(elem):
    """获取元素的文本内容"""
    text = ''
    if elem.text:
        text += elem.text
    for child in elem:
        if child.text:
            text += child.text
        if child.tail:
            text += child.tail
    return text


def create_new_table(body, rows_data):
    """创建新表格"""
    rows = len(rows_data)
    cols = len(rows_data[0]) if rows_data else 0
    
    tbl = etree.Element(f'{{{nsmap["w"]}}}tbl')
    
    # 表格属性
    tblPr = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tblPr')
    tblStyle = etree.SubElement(tblPr, f'{{{nsmap["w"]}}}tblStyle')
    tblStyle.set(f'{{{nsmap["w"]}}}val', 'TableGrid')
    tblW = etree.SubElement(tblPr, f'{{{nsmap["w"]}}}tblW')
    tblW.set(f'{{{nsmap["w"]}}}w', '0')
    tblW.set(f'{{{nsmap["w"]}}}type', 'auto')
    
    # 表格网格
    tblGrid = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tblGrid')
    for _ in range(cols):
        etree.SubElement(tblGrid, f'{{{nsmap["w"]}}}gridCol')
    
    # 表格内容
    for row_idx, row_data in enumerate(rows_data):
        tr = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tr')
        for col_idx, cell_text in enumerate(row_data):
            tc = etree.SubElement(tr, f'{{{nsmap["w"]}}}tc')
            p = etree.SubElement(tc, f'{{{nsmap["w"]}}}p')
            r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
            rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
            sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
            sz.set(f'{{{nsmap["w"]}}}val', '16')  # 8pt
            if row_idx == 0:
                b = etree.SubElement(rPr, f'{{{nsmap["w"]}}}b')
            t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
            t.text = str(cell_text)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return tbl


def main():
    print("Loading v2.10 document...")
    doc = Document(V210_PATH)
    print(f"  段落数: {len(doc.paragraphs)}")
    print(f"  表格数: {len(doc.tables)}")
    
    body = doc.element.body
    
    # 找到表格9
    tables = body.findall(f'.//{{{nsmap["w"]}}}tbl')
    table_9 = None
    table_9_idx = -1
    for i, tbl in enumerate(tables):
        text = ''
        for p in tbl.findall(f'.//{{{nsmap["w"]}}}p'):
            text += get_element_text(p)
        if '美国政府债务/GDP' in text:
            table_9 = tbl
            table_9_idx = i
            print(f"  找到表格9（索引 {i}）")
            break
    
    if table_9 is None:
        print("ERROR: 未找到表格9")
        return
    
    # 准备新表格数据
    print("\n创建新表格数据...")
    new_table_data = [
        ['指标'] + [str(y) for y in YEARS],
        ['美国政府债务/GDP'] + [f'{v}%' for v in US_GOV_DEBT],
        ['美国居民杠杆率'] + [f'{v}%' for v in US_HH_LEVERAGE],
        ['美国企业杠杆率'] + [f'{v}%' for v in US_CORP_LEVERAGE],
        ['中国政府债务/GDP'] + [f'{v}%' for v in CN_GOV_DEBT],
        ['中国居民杠杆率'] + [f'{v}%' for v in CN_HH_LEVERAGE],
        ['中国企业杠杆率'] + [f'{v}%' for v in CN_CORP_LEVERAGE],
        ['日本政府债务/GDP'] + [f'{v}%' for v in JP_GOV_DEBT],
        ['日本居民杠杆率'] + [f'{v}%' for v in JP_HH_LEVERAGE],
        ['日本企业杠杆率'] + [f'{v}%' for v in JP_CORP_LEVERAGE],
    ]
    
    # 创建新表格
    print("创建新表格...")
    new_tbl = create_new_table(body, new_table_data)
    
    # 替换旧表格
    print("替换旧表格...")
    table_9.addprevious(new_tbl)
    body.remove(table_9)
    
    # 保存
    print(f"\nSaving to {OUT_PATH}...")
    doc.save(OUT_PATH)
    
    # 验证
    verify_doc = Document(OUT_PATH)
    print(f"\nVerification:")
    print(f"  段落数: {len(verify_doc.paragraphs)}")
    print(f"  表格数: {len(verify_doc.tables)}")
    
    # 验证表格9
    for i, table in enumerate(verify_doc.tables):
        if i == 8:
            print(f"\n  表格9更新后:")
            print(f"    行数: {len(table.rows)}")
            print(f"    列数: {len(table.columns)}")
            # 显示前3行
            for r, row in enumerate(table.rows[:3]):
                cells = [cell.text[:10] for cell in row.cells[:5]]
                print(f"    行{r}: {cells}...")
    
    print(f"  ✅ Done!")


if __name__ == "__main__":
    main()
