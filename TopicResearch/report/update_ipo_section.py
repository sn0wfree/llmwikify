"""
更新 1.3.4 巨型IPO现象章节：添加AI IPO抽血分析
"""
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

DOCX_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.5.docx'
OUTPUT_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'

def add_paragraph_after(doc, ref_text, text, font_size=9, bold=False):
    """在指定段落后添加新段落"""
    for i, p in enumerate(doc.paragraphs):
        if ref_text in p.text:
            # 在下一个段落前插入（通过在末尾添加后移动）
            new_p = doc.add_paragraph()
            run = new_p.add_run(text)
            run.font.size = Pt(font_size)
            run.font.bold = bold
            return True
    return False

def add_table_after(doc, ref_text, headers, rows):
    """在指定段落后添加表格"""
    for i, p in enumerate(doc.paragraphs):
        if ref_text in p.text:
            # 创建表格
            table = doc.add_table(rows=len(rows)+1, cols=len(headers))
            # 使用默认样式
            
            # 表头
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(8)
            
            # 数据行
            for r, row_data in enumerate(rows):
                for c, val in enumerate(row_data):
                    cell = table.rows[r+1].cells[c]
                    cell.text = str(val)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            return True
    return False

def main():
    print('Loading document...')
    doc = docx.Document(DOCX_PATH)
    
    # 找到1.3.4节位置并添加内容
    # 注意：python-docx只能在末尾添加，所以我们在末尾添加后标记位置
    
    print('\nAdding AI IPO content...')
    
    # 添加分页符
    doc.add_page_break()
    
    # 添加补充章节
    p = doc.add_heading('补充：本轮AI巨型IPO抽血分析', level=2)
    
    # 添加背景说明
    p = doc.add_paragraph()
    run = p.add_run(
        '2026年6月，SpaceX在纳斯达克成功上市（股票代码SPCX），募资750亿美元，创下全球资本市场史上最大IPO纪录。'
        '紧随其后，OpenAI计划于2026Q4-2027Q1进行IPO，规模预计1000亿美元。'
        '本轮AI巨兽集中上市的"抽血效应"值得高度关注。'
    )
    run.font.size = Pt(9)
    
    # 表格1：本轮AI IPO全景
    p = doc.add_paragraph()
    run = p.add_run('表 1.3-1：本轮AI巨型IPO全景')
    run.font.size = Pt(9)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers = ['公司', '状态', '募资规模', '估值', '抽血占比']
    rows = [
        ['SpaceX', '✅ 已上市 (2026.6.12)', '$750亿', '$1.77万亿', '0.17%'],
        ['OpenAI', '拟IPO (2026Q4-2027Q1)', '$1000亿', '$8500亿', '0.22%'],
        ['Anthropic', '拟IPO (2027+)', '$500亿', '$2000亿', '0.11%'],
        ['合计', '-', '$2250亿', '-', '0.50%'],
    ]
    add_table_after(doc, '本轮AI巨兽集中上市', headers, rows)
    
    # 表格2：SpaceX vs 中石油 IPO对比
    p = doc.add_paragraph()
    run = p.add_run('表 1.3-2：SpaceX vs 中石油 IPO对比')
    run.font.size = Pt(9)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    headers2 = ['维度', '2007 中石油', '2026 SpaceX', '差异']
    rows2 = [
        ['募资规模', '¥668亿 ($97亿)', '$750亿', 'SpaceX是中石油7.7倍'],
        ['上市估值', '$1.08万亿', '$1.77万亿', 'SpaceX高64%'],
        ['首日涨幅', '+191% (开盘)', '+29% (开盘)', '中石油更疯狂'],
        ['冻结资金', '¥3.3万亿 (40%)', '$2500亿认购 (4倍)', '中石油抽血更猛'],
        ['散户参与', '>60%', '30%', '散户占比下降'],
        ['盈利状态', '盈利', '亏损 ($49.4亿)', 'SpaceX尚未盈利'],
        ['后续表现', '一年内破发', '待观察', '-'],
    ]
    add_table_after(doc, '本轮AI巨兽集中上市', headers2, rows2)
    
    # 添加抽血机制分析
    p = doc.add_paragraph()
    run = p.add_run('1.3.4-AI IPO抽血机制分析')
    run.font.size = Pt(10)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run(
        '（1）抽血规模对比：SpaceX募资750亿美元，占标普500流通市值约0.17%；'
        'OpenAI预计募资1000亿美元，占纳斯达克100流通市值约0.4%。'
        '合计潜在抽血2250亿美元，占标普500流通市值约0.50%。'
        '对比2007年中石油冻结3.3万亿元（占A股流通市值40%），本轮抽血规模远小于2007年。'
    )
    run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run(
        '（2）存量占用更严重：本轮AI牛市中，NVIDIA+Microsoft已占纳斯达克100约24%权重。'
        'SpaceX上市后迅速纳入纳斯达克100，进一步加剧头部集中。'
        '当AI巨兽占据指数主要权重时，新IPO的边际抽血效应可能被放大。'
    )
    run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run(
        '（3）流动性环境差异：2007年M2增速17%+外汇占款被动投放，流动性极度宽松；'
        '2026年M2增速8%+美联储维持高利率，流动性相对紧张。'
        '这意味着本轮IPO对市场流动性的冲击可能更敏感。'
    )
    run.font.size = Pt(9)
    
    p = doc.add_paragraph()
    run = p.add_run(
        '（4）散户参与度下降：SpaceX散户份额30%（远高于近年平均水平），但中石油时代散户占比超60%。'
        '机构主导的IPO通常波动更小，但一旦下跌，机构止损盘可能更猛烈。'
    )
    run.font.size = Pt(9)
    
    # 添加关键判断
    p = doc.add_paragraph()
    run = p.add_run('关键判断：')
    run.font.size = Pt(9)
    run.font.bold = True
    
    p = doc.add_paragraph()
    run = p.add_run(
        '本轮AI IPO抽血规模虽小于2007年，但需高度关注两个风险点：'
        '（1）多只AI巨兽集中上市的时间窗口重叠（SpaceX已上市+OpenAI即将IPO）；'
        '（2）存量占用严重（NVIDIA+Microsoft+SpaceX已占纳斯达克100约26%）。'
        '核心监测指标：OpenAI IPO后30日表现（S2信号触发点）。'
    )
    run.font.size = Pt(9)
    
    # 保存
    print(f'\nSaving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'✅ Done! Saved to {OUTPUT_PATH}')

if __name__ == '__main__':
    main()
