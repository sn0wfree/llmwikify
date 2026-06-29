"""
添加新章节 1.1.6 蒙代尔不可能三角
"""
import os
import copy
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.expanduser("~/llmwikify/TopicResearch/report")
CHARTS_DIR = os.path.join(BASE_DIR, "charts")
V28_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.8.docx")
OUT_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.9.docx")

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def get_element_text(elem):
    """获取元素的文本内容"""
    text = ''
    if elem.text:
        text += elem.text
    for child in elem:
        if child.text:
            text += child.text
        if child.tail:
            text += child.tail
    return text


def find_element_index(body, keyword):
    """在 body 中找到包含关键词的段落索引"""
    for i, elem in enumerate(body):
        if elem.tag == f'{{{nsmap["w"]}}}p':
            text = get_element_text(elem)
            if keyword in text:
                return i
    return -1


def create_paragraph_element(text, bold=False, size=9, alignment=None):
    """创建段落 XML 元素"""
    p = etree.Element(f'{{{nsmap["w"]}}}p')
    
    pPr = etree.SubElement(p, f'{{{nsmap["w"]}}}pPr')
    
    if alignment:
        jc = etree.SubElement(pPr, f'{{{nsmap["w"]}}}jc')
        jc.set(f'{{{nsmap["w"]}}}val', alignment)
    
    r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
    
    rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{nsmap["w"]}}}rFonts')
    rFonts.set(f'{{{nsmap["w"]}}}ascii', '微软雅黑')
    rFonts.set(f'{{{nsmap["w"]}}}hAnsi', '微软雅黑')
    
    sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
    sz.set(f'{{{nsmap["w"]}}}val', str(size * 2))
    
    if bold:
        b = etree.SubElement(rPr, f'{{{nsmap["w"]}}}b')
    
    t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return p


def create_table_element(table_data):
    """创建表格 XML 元素"""
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    
    tbl = etree.Element(f'{{{nsmap["w"]}}}tbl')
    
    tblPr = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tblPr')
    tblStyle = etree.SubElement(tblPr, f'{{{nsmap["w"]}}}tblStyle')
    tblStyle.set(f'{{{nsmap["w"]}}}val', 'TableGrid')
    tblW = etree.SubElement(tblPr, f'{{{nsmap["w"]}}}tblW')
    tblW.set(f'{{{nsmap["w"]}}}w', '0')
    tblW.set(f'{{{nsmap["w"]}}}type', 'auto')
    
    tblGrid = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tblGrid')
    for _ in range(cols):
        etree.SubElement(tblGrid, f'{{{nsmap["w"]}}}gridCol')
    
    for row_idx, row_data in enumerate(table_data):
        tr = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tr')
        for col_idx, cell_text in enumerate(row_data):
            tc = etree.SubElement(tr, f'{{{nsmap["w"]}}}tc')
            p = etree.SubElement(tc, f'{{{nsmap["w"]}}}p')
            r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
            rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
            sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
            sz.set(f'{{{nsmap["w"]}}}val', '18')
            if row_idx == 0:
                b = etree.SubElement(rPr, f'{{{nsmap["w"]}}}b')
            t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
            t.text = str(cell_text)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return tbl


def create_impossible_triangle_section():
    """创建蒙代尔不可能三角章节内容"""
    elements = []
    
    # 标题
    elements.append(create_paragraph_element(
        "1.1.6 蒙代尔不可能三角与中国政策选择",
        size=10, bold=True))
    
    # 说明文字
    elements.append(create_paragraph_element(
        '蒙代尔-弗莱明不可能三角（Mundell-Fleming Trilemma）是国际金融学的经典理论：'
        '固定汇率、资本自由流动、独立货币政策三者不可兼得，任何国家最多只能同时实现其中两个目标。',
        size=9))
    
    elements.append(create_paragraph_element(
        '该理论对于理解2005-2007年中国货币政策困境具有关键意义。彼时中国选择的是"B组合"——'
        '固定汇率（盯住美元8.28）+ 资本管制（名义上禁止热钱流入），代价是牺牲货币政策独立性，'
        '央行被迫通过大规模冲销操作来维持流动性平衡。',
        size=9))
    
    # 表格标题
    elements.append(create_paragraph_element(
        "表 1.1-3：蒙代尔不可能三角下中美政策选择对比",
        size=9, bold=True, alignment='center'))
    
    # 表格
    elements.append(create_table_element([
        ['维度', '中国 05-07', '中国 现在', '美国'],
        ['汇率制度', '软盯住美元（锁定8.28）', '管理浮动（±2%波动）', '自由浮动'],
        ['资本流动', '名义管制（热钱绕道流入）', '渐进开放（防外流为主）', '完全自由'],
        ['货币政策', '受冲销约束（被动宽松）', '独立性增强（主动宽松）', '完全独立'],
        ['三角位置', 'B组合（左上角）', '趋向C组合（中间偏右下）', 'C组合（右下角）'],
        ['外储/GDP', '15%→37%（快速积累）', '17%（相对稳定）', '0.3%（无需外储）'],
        ['主要矛盾', '升值压力+热钱流入', '贬值压力+资本外流', '通胀 vs 就业'],
    ]))
    
    # 补充说明
    elements.append(create_paragraph_element(
        '2005-2007年，中国处于三角形左上角的B组合位置：'
        '通过固定汇率（8.28）和资本管制（名义上）来维持出口竞争力，'
        '但代价是央行必须被动对冲外汇占款，冲销成本约占GDP的1-2%。'
        '加息→热钱流入→更多冲销→成本更高，形成自我强化循环。',
        size=9))
    
    elements.append(create_paragraph_element(
        '2015年"811汇改"后，中国逐步向C组合过渡：'
        '汇率转向管理浮动（±2%波动），资本管制从"防流入"转向"防流出"，'
        '货币政策独立性显著增强。外储从4万亿美元峰值降至3.2万亿，'
        '反映了央行主动放弃部分外储以换取政策空间。',
        size=9))
    
    elements.append(create_paragraph_element(
        '美国始终处于C组合位置：自由浮动汇率+完全资本流动+完全独立货币政策。'
        '美联储无需积累外汇储备，政策目标单一（通胀+就业），'
        '这也是美元作为全球储备货币的制度优势所在。',
        size=9))
    
    return elements


def main():
    print("Loading v2.8 document...")
    doc = Document(V28_PATH)
    print(f"  段落数: {len(doc.paragraphs)}")
    print(f"  表格数: {len(doc.tables)}")
    
    body = doc.element.body
    
    # 找到 1.1.5 财政政策 位置
    s115_idx = find_element_index(body, "1.1.5 财政政策")
    if s115_idx == -1:
        print("ERROR: Cannot find section 1.1.5")
        return
    
    print(f"  1.1.5 section at element: {s115_idx}")
    
    # 找到 1.1.6 资本市场 位置
    s116_idx = find_element_index(body, "1.1.6 资本市场")
    if s116_idx == -1:
        print("ERROR: Cannot find section 1.1.6")
        return
    
    print(f"  1.1.6 section at element: {s116_idx}")
    
    # 插入新章节
    print("\nInserting new section 1.1.6...")
    new_section = create_impossible_triangle_section()
    
    for i, elem in enumerate(new_section):
        body.insert(s116_idx + i, elem)
    
    print(f"  Inserted {len(new_section)} elements")
    
    # 保存
    print(f"\nSaving to {OUT_PATH}...")
    doc.save(OUT_PATH)
    
    # 验证
    verify_doc = Document(OUT_PATH)
    print(f"\nVerification:")
    print(f"  段落数: {len(verify_doc.paragraphs)}")
    print(f"  表格数: {len(verify_doc.tables)}")
    
    print(f"  ✅ Done!")


if __name__ == "__main__":
    main()
