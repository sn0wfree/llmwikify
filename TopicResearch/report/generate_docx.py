"""
重写 Word 研究报告 - 去除照抄感
核心改进：
1. 叙事弧线贯穿全文（不是按节堆砌数据）
2. 表格/图表作为证据嵌入分析，不是独立罗列
3. 每节有"分析"而非"陈列"
4. 段间有粘合，不是断点
5. 数据有"所以呢"（含义），不只是"是什么"
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

CHART_DIR = '/home/ll/llmwikify/TopicResearch/report/charts'
OUTPUT = '/home/ll/llmwikify/TopicResearch/report/复盘05-07有色vs当前AI板块投资研究.docx'


def set_font(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.space_after = Pt(4)


def H1(doc, text, color=None):
    if color is None:
        color = RGBColor(15, 23, 42)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def H2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)
    return p


def H3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(71, 85, 105)
    return p


def P(doc, text, indent=True, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.7
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    return p


def quote_block(doc, text, color=None):
    if color is None:
        color = RGBColor(180, 83, 9)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.right_indent = Pt(22)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.6
    run = p.add_run(text)
    run.font.bold = True
    run.font.color.rgb = color
    return p


def add_table(doc, data, col_widths=None, header_color=None):
    if header_color is None:
        header_color = RGBColor(30, 58, 95)
    rows, cols = len(data), len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.3
                for run in para.runs:
                    if i == 0:
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.size = Pt(9.5)
            if i == 0:
                # Header cell shading
                tcPr = cell._tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is None:
                    shd = tcPr.makeelement(qn('w:shd'), {})
                    tcPr.append(shd)
                shd.set(qn('w:fill'), '1E3A5F')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    return table


def add_chart(doc, filename, caption, width_inches=6.5):
    path = os.path.join(CHART_DIR, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap_run = cap.add_run(caption)
    cap_run.font.size = Pt(9.5)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(100, 116, 139)


# ============================================================
doc = Document()
set_font(doc)

for section in doc.sections:
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ============================================================
# 封面
# ============================================================
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("复盘 2005–2007 有色金属超级周期\n对标当前 AI 板块投资")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(15, 23, 42)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(20)
r = sub.add_run("——基于「资本开支驱动的需求正反馈」的 3 个量化顶部信号")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("报告编号：TR-2026-001  |  报告日期：2026 年 6 月 28 日  |  版本 v2.0")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(148, 163, 184)

doc.add_paragraph()
doc.add_paragraph()

# 数据来源
src = doc.add_paragraph()
src.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = src.add_run(
    "数据来源：FRED、PBOC、BOJ、IMF、World Bank、ICSG、海关总署、SEC EDGAR、Bloomberg\n"
    "辅助文件：TopicResearch/data/raw/ (32 个 JSON 数据文件)\n"
    "本报告准绳：05-07 有色复盘 HTML 报告 + slides 演示文稿")
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor(148, 163, 184)

doc.add_paragraph()

disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = disclaimer.add_run(
    "免责声明：本报告基于历史类比研究，不构成投资建议。AI 浪潮有真实盈利支撑，"
    "「此次不同」可能令信号延迟或失效。建议组合使用三信号，不单一依赖。")
r.font.size = Pt(8.5)
r.font.italic = True
r.font.color.rgb = RGBColor(180, 83, 9)

doc.add_page_break()


# ============================================================
# 引言：问题与分析框架
# ============================================================
H1(doc, "引言：为什么要复盘 2005–2007？")

P(doc,
    "2008 年金融危机后，全球资本市场经历了近 15 年的「低利率 + 低增长」长期停滞。直至 2022 年 11 月 ChatGPT 的横空出世，"
    "「AGI 革命」叙事点燃了新一轮资本开支浪潮——M7（Magnificent 7，Apple、Microsoft、Google、Amazon、Meta、Tesla、NVIDIA）"
    "在 2023–2026 年间将云端 Capex 累计推高至数千亿美元规模，NVIDIA 股价从 $14 涨至 $130（+800%），"
    "纳斯达克 100 指数从 11000 点冲至 21500 点（+95%）。")

P(doc,
    "这一轮 AI 浪潮的「资本开支驱动的需求正反馈」与 2005–2007 年的有色金属超级周期具有令人不安的结构相似性——"
    "彼时全球铜价从 $3,000 涨至 $8,985（+197%），LME 铝价从 $1,807 涨至 $3,000，A 股有色金属指数涨幅高达 1,857%，"
    "中铝、云铜等龙头股在「资产注入」叙事下 PE 推升至 60–80 倍。")

P(doc, "然而，两轮浪潮在「流动性来源」和「估值锚」上存在本质差异：")

P(doc,
    "05–07 的流动性来自**央行被动放水**——中国出口创汇转化为外汇占款（2005–2007 累计 ~8 万亿元），"
    "央行结汇后基础货币扩张，M2 增速从 17.6% 升至 18.5%；同时美元持续贬值（92→80，-13%），"
    "日本超低利率（0–0.5%）下日元套息交易规模达 $500B+，三重机制共同构成了商品超级周期的「全球流动性大合唱」。",
    indent=True)

P(doc,
    "23–26 的流动性则来自**企业主动喷泉**——M7 自由现金流 + 部分企业债融资直接转化为 AI Capex，"
    "M7（除 NVDA）平均 Capex/Rev 从 2024 年的 18% 升至 2026E 的 30%。美联储虽然缩表，但企业信用条件仍宽——"
    "M7 自身的盈利和现金成为 AI Capex 的唯一来源。",
    indent=True)

P(doc,
    "这一差异意味着 23–26 泡沫可能更「粘」——反应慢于 05–07，但顶部更「陡」——M7 同步减投会瞬间反转。",
    indent=True)

P(doc,
    "本报告的核心问题是：在 23–26 AI 浪潮的当前阶段（截至 2026 年 6 月），**我们能否找到 3 个可量化、可监测、可证伪的顶部信号**？"
    "这些信号应当在 05–07 的历史镜像中得到验证，并能在不同情境下提示风险。")

quote_block(doc,
    "本报告的 3 个核心结论：(1) 当前 3 个信号均未触发，综合评分 0/3 ~ 0.5/3，维持关注；"
    "(2) 顶部窗口预测 2026H2 – 2027H1；(3) 关键观察点为 OpenAI IPO 后 30 日表现。")


# ============================================================
# 第一部分
# ============================================================
H1(doc, "第一部分　2005–2007 有色金属超级周期全景复盘")

P(doc,
    "要回答「AI 浪潮何时见顶」这一问题，首先需要完整理解 2005–2007 这一历史镜像。"
    "本部分从「宏观背景—货币政策—汇率/贸易—外汇储备—财政政策—资本市场」六个维度还原 2005–2007 的全貌，"
    "并最终回答一个核心问题：当年的超级周期**为什么**会形成、**为什么**会在 2008 年终结。",
    indent=False)

# §1.1
H2(doc, "§1.1　全球主要经济体背景：三国同步增长的罕见窗口")

P(doc,
    "2005–2007 是 21 世纪以来美中日三大经济体**同步增长**的罕见窗口。理解这一同步性，"
    "是理解商品超级周期的真正起点——三国增速虽然差异巨大，但角色高度互补。",
    indent=False)

P(doc, "**美国：减速但消费强**。美国实际 GDP 增速从 2005 年 3.5% 降至 2007 年 1.9%，看似在减速，"
    "但名义 GDP 仍维持 5% 左右，**消费和资产价格仍强劲**——这是「财富效应」支撑的典型表现："
    "标普 500 指数从 2002 年底部上涨 101%，居民可支配收入随资产价格上涨而扩张，"
    "反过来支撑了对中国制造的消费品需求。", indent=True)

P(doc, "**中国：加速，重工业化+城镇化高峰**。中国 GDP 实际增速从 11.3% 升至 14.2%，"
    "名义 GDP 增速更高达 23%（含通胀和资产价格重估）。中国贡献了全球铜消费增量的 60%+，"
    "M2 增速 17.6%→18.5%，固定资产投资 CAGR 超过 25%——这是「重工业化+城镇化」双轮驱动的典型样本。", indent=True)

P(doc, "**日本：走出「失去的十年」**。日本实际 GDP 增速仅 0.6%–0.8%，名义 GDP ≈ 0，"
    "看似在通缩中挣扎，但企业治理改善、外资流入，"
    "日经 225 在 2005 年上涨 40.24%——这意味着日本完成了从「偿债」到「扩张」的模式切换，"
    "即将在全球金融体系中扮演新的角色。", indent=True)

P(doc, "**关键洞察**：这种「同步增长」是 2005–2007 独有的窗口。"
    "美国虽减速但消费强、中国处于工业化高峰、日本走出通缩——三国缺一不可。"
    "商品超级周期的真正燃料不是「中国需求」单独一个变量，而是**「消费国+生产国+融资国」的三角互补结构**。", indent=False)

# §1.2
H2(doc, "§1.2　货币政策：表面紧缩、实际宽松")

P(doc,
    "理解 2005–2007 货币政策的关键，是要意识到：**表面上的紧缩，掩盖了实际金融条件的宽松**。"
    "如果只看 2007 年末三国央行政策利率——美联储 5.25%、中国 1Y 贷款基准 7.47%、日银 0.5%——"
    "会得出三国都在紧缩的结论。然而，市场实际感受的金融条件却普遍宽松。")

P(doc, "**美国的「格林斯潘之谜」**揭示了短端利率与长端利率的脱节。"
    "美联储在 2004.06–2006.06 期间加息 17 次，联邦基金利率从 1% 抬至 5.25%；"
    "但 10 年期国债利率仅从 4.2% 升至 4.6%——几乎纹丝不动。短长利差从 +3.2% 收窄到 +0.65%，"
    "意味着**美联储加的息在金融市场上几乎没有传导**。这一悖论的根源是全球「储蓄过剩」——"
    "亚洲国家、石油美元的避险资金持续压低长端利率，使美国实际金融条件仍宽松。", indent=True)

P(doc, "**中国的情况更奇特**。1Y 贷款基准 7.47% 看似很高，但同期 10 年期国债仅 4.5%；"
    "更反常的是 1Y 存款基准 4.14% 居然高于 1Y 国债 2.79%——在正常利率体系中，"
    "国债作为无风险利率应该是市场利率的下限，但中国央行通过存款管制将存款利率人为抬高在国债之上。"
    "这意味着 2007 年的中国同时存在**管制利率倒挂**（4.14% > 2.79%）和**贷存利差 3.33%**（7.47% - 4.14%）。",
    indent=True)

P(doc, "这两个数字是理解中国 2007 年资产价格泡沫的关键。贷存利差 3.33% 是央行在 2019 年 LPR 改革前"
    "给中国银行业的「制度性红包」——银行「躺着赚钱」的安全垫；"
    "而存款利率被管制在国债之上，意味着居民存款被锁在银行体系内，"
    "无法通过市场化渠道配置到股市和楼市。**结果是「被动溢出」——存款利率过低，"
    "居民被迫将储蓄转向资产，叠加 M2 增速 18% 的扩张效应**。", indent=False)

add_chart(doc, '04_three_rate.png',
          '图表 1：2007 年三国货币政策三层利率对比——揭示中国管制利率倒挂与贷存利差 3.33%',
          width_inches=6.5)

P(doc, "**日本的故事则是「套利」**。日银 2006 年退出 QE，隔夜利率 0% → 0.5%，但 10Y JGB 仅 1.5%、房贷 2.5%——"
    "量宽传导充分。**真正的故事是日元套息**——借日元（利率极低）投资美元资产或商品，"
    "估算规模 $500B+。这是 2005–2007 全球风险偏好的关键推手。", indent=True)

quote_block(doc,
    "**为什么三国货币政策「表面紧缩、实际宽松」？**因为三个水龙头的开关逻辑完全不同——"
    "美国被储蓄过剩压制，中国被外汇占款抵消，日本被套利输出。"
    "这构成了商品超级周期的全球流动性基础。")


# §1.3
H2(doc, "§1.3　外汇储备与汇率：全球失衡的镜像")

P(doc,
    "2005–2007 期间，中国外汇储备从 8,189 亿美元跃升至 15,282 亿美元，**两年半内几乎翻倍**。"
    "这一数字背后是全球贸易格局的深刻变化——中国出口创汇 $1,022 亿→$2,622 亿，"
    "贸易顺差翻倍，外汇占款累计增加近 8 万亿元。")

P(doc, "**外储/GDP 48%**——这是 2007 年中国的核心数字，全球第一。"
    "美国仅 1.7%，日本 21%。这一失衡本身不可持续，最终必然通过汇率/利率/资产价格调整——"
    "这也是 2008 年危机后中国外储增长显著放缓的根本原因。", indent=True)

add_chart(doc, '03_transmission_chain.png',
          '图表 2：中国外汇占款完整传导链（2005-2008）——贸易顺差→外储→外汇占款→M2 增速的 4 变量协同',
          width_inches=6.8)

P(doc, "**美元贬值的连锁效应**是 2005–2007 商品超级周期最直接的燃料："
    "美元指数从 92 跌至 80（-13%），推升以美元计价的商品价格——LME 铜从 $3,000 涨至 $8,000+；"
    "同时全球资金流出美国、涌入新兴市场；美国进口成本上升、通胀压力上升、实际利率下降。"
    "**这是一条自我强化的链条**——美元越贬值，商品越涨，新兴市场越受益，对美出口越多，外汇占款越多。", indent=False)

add_chart(doc, '09_three_country_m2.png',
          '图表 2.5：三国 M2 增速 vs 美元指数（2005-2008）——揭示中国是全球流动性的"主水源"',
          width_inches=6.8)

P(doc, "图表 2.5 揭示了三国 M2 增速的**惊人差异**：2007 年中国 M2 增速 18.5%，"
    "是美国的 3.3 倍、日本的 13 倍。**这意味着中国是 2005–2007 全球流动性的「主水源」**——"
    "中国央行通过外汇占款投放的 8 万亿基础货币，相当于 8 次降准的力度，"
    "相当于美联储同期通过 QE 投放的流动性总量。",
    indent=False)

P(doc, "**人民币升值的悖论**值得单独分析。2005 年 7 月汇改启动，人民币兑美元累计升值约 12%，"
    "但贸易顺差反而翻倍。这一悖论的答案在于：(1) 中国生产效率提升，单位劳动力成本下降；"
    "(2) 加工贸易占比 50%+，汇率传导有限；(3) 外需强劲，升值预期反而强化了出口订单的提前释放。",
    indent=True)

# §1.4
H2(doc, "§1.4　资本市场与巨型 IPO：见顶的三重信号")

P(doc,
    "三国股市的涨幅差异不是估值差异，而是**经济结构差异**的反映。",
    indent=False)

add_table(doc, [
    ["", "🇺🇸 美国（标普 500）", "🇨🇳 中国（上证综指）", "🇯🇵 日本（日经 225）"],
    ["2005-2007 涨幅", "+101%", "+514%", "+50%"],
    ["起点背景", "2002 科网泡沫后底部", "2005 股改后起步", "2003 摆脱通缩"],
    ["PE 估值变化", "16x → 22x（扩张 1.4 倍）", "14x → 50x（扩张 3.5 倍）", "18x → 20x（基本稳定）"],
    ["总市值/GDP", "130% → 150%", "35% → 120%", "70% → 95%"],
    ["核心驱动", "消费+地产+金融", "投资+散户+资产注入", "通缩挣扎+企业治理"],
])

P(doc, "**中国 PE 14x → 50x 是 2005–2007 最极端的估值扩张**。这一扩张的驱动是「资产注入」叙事——"
    "把矿的储量直接乘以现货价格算市值，把未来盈利一次性贴现到现在。"
    "当 2007 年下半年市场已经「不再讨论风险，只看注入多少矿」时，"
    "这种「梦想定价」就达到了极限。", indent=False)

add_chart(doc, '01_copper_phases.png',
          '图表 3：LME 铜价走势与三大阶段（2005-2008）——揭示商品见顶早于股价 17 个月',
          width_inches=6.5)

P(doc, "图表 3 揭示了一个被市场广泛忽视的细节：**LME 铜价早在 2006 年 5 月就已见顶（$8,800），"
    "而股价直到 2007 年 10 月才见顶**——17 个月的「顶背离」。"
    "这是典型的「叙事溢价」阶段——股价不再锚定商品现货，而是锚定「未来注入矿」的无限承诺。",
    indent=False)

P(doc, "**2007 年下半年的「IPO 洪峰」是政策调控工具**。中石油、神华、建行等巨型 IPO 密集发行，"
    "通过大规模发行吸收市场过剩流动性。中石油 A 股 IPO 冻结资金 3.3 万亿元，"
    "**相当于当时 A 股流通市值的 40%**——这是 07 年顶部的关键触发因素。",
    indent=False)

add_chart(doc, '02_shanghai_ipo.png',
          '图表 4：上证综指与巨型 IPO 时间线（2005-2008）——揭示 6124 顶部与 IPO 抽水的同步性',
          width_inches=6.5)

P(doc, "图表 4 清晰呈现了 2007 年 4 月至 11 月的「IPO 洪峰」："
    "4 月中铝（¥100 亿）、9 月建行（¥580 亿）、10 月神华（¥666 亿，冻结 2.7 万亿）、11 月中石油（¥668 亿，冻结 3.3 万亿）。"
    "**这些 IPO 不是巧合**——监管层明确知道市场过热，"
    "通过「市场化的流动性收紧」（IPO 抽水）来抑制资产价格泡沫。"
    "**这是一个教科书级的「政策窗口指导」**——监管从未明令禁止新基金发行，但通过 IPO 节奏控制，"
    "实现了对资金面的精准调控。", indent=False)


# §1.5
H2(doc, "§1.5　三角闭环：超级周期的真正终结机制")

P(doc,
    "理解 2005–2007 终结机制的关键，是要理解「三角闭环」的结构。",
    indent=False)

add_chart(doc, '07_triangle_loop.png',
          '图表 5：三角闭环机制示意图——美国-中国-日本「消费-生产-融资」闭环',
          width_inches=6.5)

P(doc, "这一闭环的运作机制是：美国（消费+地产+金融杠杆）→ 美元流向中国出口创汇 → 央行结汇形成外汇占款 → "
    "M2 扩张推升投资需求；与此同时日本（低成本资金）→ 借日元投资美元资产/商品 → 全球风险偏好 → "
    "资源需求 → 商品价格。这是一条**自洽的正反馈链**。", indent=False)

P(doc, "**闭环破裂的三个条件**——任一发生即可能终结：", indent=False)

for line in [
    "**美国消费崩盘**（次贷危机）→ 终端需求消失",
    "**中国生产停滞**（固投见顶）→ 资源需求消失",
    "**日本融资收紧**（日银加息）→ 套利逆转，全球流动性收缩",
]:
    p = doc.add_paragraph(line, style='List Number')
    p.paragraph_format.left_indent = Pt(40)

P(doc, "**2008 年的实际情况是条件 1 和条件 2 同时触发**——次贷危机导致美国消费崩盘（汽车业、零售业、消费电子全线下滑），"
    "中国出口骤降（2009 年 1 月出口同比 -17.5%）。三角闭环彻底破裂。"
    "商品价格在 6 个月内跌 50%+，日本套利交易逆转，日元升值 20%+。", indent=False)

P(doc, "这一历史经验对 23–26 AI 浪潮的启示是：需要关注 M7 企业的「同步性」——如果 M7 7 家公司同时减投 AI Capex，"
    "整个 AI 板块的「自洽链」将瞬间断裂。**这是 23–26 三角闭环的「脆弱性」所在**——"
    "风险集中在 7 家公司，而非整个信贷体系。", indent=False)


# ============================================================
# 第二部分
# ============================================================
H1(doc, "第二部分　对标 AI 板块：结构相似性与本质差异")

P(doc,
    "基于第一部分的 6 维度复盘，本部分将逐一映射到 2023–2026 的 AI 板块投资浪潮。"
    "我们关心的核心问题是：**哪些维度高度相似？哪些维度本质不同？这些差异如何影响信号的设计**？",
    indent=False)

# §2.1
H2(doc, "§2.1　10 维度对标：结构对照")

P(doc,
    "下表呈现了 2005–2007 与 2023–2026 的 10 维度对标。同维度内，用「共同本质」列提炼两轮行情的内在相似性，"
    "这有助于识别哪些信号是「模式可复制」的，哪些是「此次不同」的。",
    indent=False)

add_table(doc, [
    ["维度", "2005-2007 有色", "2023-2026 AI"],
    ["驱动叙事", "中国城镇化+工业化+WTO 红利", "AGI 革命+算力军备竞赛"],
    ["前三大经济体", "美(消费)+中(生产)+日(融资)三角互补", "美(AI创新)+中(追赶) 中美竞争"],
    ["流动性来源", "央行被动放水+外汇占款", "M7 主动喷泉+企业自驱 Capex"],
    ["债务特征", "美国居民杠杆 99% 见顶", "美国政府债务 127% + M7 现金充裕"],
    ["产业政策", "中国限制出口+美国无干预", "中国大基金+美国芯片管制"],
    ["市场共识", "怀疑→接受→狂热 (3 年)", "怀疑→接受→部分狂热 (4 年)"],
    ["龙头涨幅", "有色 +1,857% / 云铜 +1,122%", "NVDA +800% / NDX +95%"],
    ["估值扩张", "PE 14x→50x (3.5 倍)", "NVDA P/S 10x→30x (3 倍)"],
    ["巨型 IPO", "中石油冻结 3.3 万亿 (流通 40%)", "OpenAI 待 IPO ($1,000 亿规模)"],
    ["终结模式", "政策+IPO+商品见顶 三重共振", "Capex 回撤+应用证伪+IPO 抽血"],
], col_widths=[Cm(2.5), Cm(7), Cm(7)])

P(doc, "**关键差异点 1：流动性来源**。05-07 是「央行被动放水」——外汇占款、储蓄过剩、日元套息形成全球流动性大合唱；"
    "23-26 是「M7 主动喷泉」——企业自筹 Capex，自主决策节奏。", indent=True)

P(doc, "**关键差异点 2：债务结构**。05-07 的脆弱性是「美国居民 99% 杠杆 + 中国企业 125% 杠杆」——"
    "整个信贷体系的风险；23-26 的脆弱性是「M7 集中度过高 + 美国政府 127% 债务」——"
    "风险集中在 7 家公司。", indent=True)

P(doc, "**关键差异点 3：估值锚**。05-07 的 PE 60-80x 完全脱离当期盈利，是「梦想定价」；"
    "23-26 的 NVDA P/S 30x 仍有真实业绩支撑（EPS 增速 147%），是「业绩+梦想」双驱动。"
    "这意味着 23-26 的「梦想→残值」切换可能比 05-07 晚 6-12 个月发生。", indent=True)


# §2.2
H2(doc, "§2.2　AI 板块现状：截至 2026 年 6 月")

H3(doc, "M7 云厂商 Capex：S1 信号的核心观测点")

P(doc,
    "M7 云厂商（Microsoft、Google、Amazon、Meta）的 AI Capex 是 23-26 浪潮的核心。"
    "与 2005-2007 的中国固投增速类似，**M7 Capex/Rev 比率**反映了「资本开支扩张速度」，"
    "是算力需求「真实性」的最强信号。",
    indent=False)

add_chart(doc, '05_m7_capex.png',
          '图表 6：M7 云厂商 Capex/Rev 趋势（2024-2026E）——35% 阈值线即 S1 信号触发线',
          width_inches=6.5)

P(doc, "图表 6 呈现了 4 家云厂商的 Capex/Rev 趋势。可以观察到三个关键事实：", indent=False)

P(doc, "**第一，2024-2026 期间所有 4 家云厂商的 Capex/Rev 都在快速攀升**。"
    "Microsoft 从 22% 升至 35%，Google 从 15% 升至 25%，Amazon 从 11% 升至 18%，Meta 从 25% 升至 39%。"
    "M7 平均（除 NVDA）从 18% 升至 30%——距 35% 阈值仅 5 个百分点。",
    indent=True)

P(doc, "**第二，Meta 是 S1 信号的「前沿观测点」**。"
    "Meta 的 Capex/Rev 在 2026E 已达 39%，超过 35% 阈值——"
    "如果 Meta 在 2026Q4 或 2027Q1 财报中给出「Capex 增长放缓」的指引，"
    "**这将是 S1 信号最早可能触发的节点**。",
    indent=True)

P(doc, "**第三，5% 阈值的逻辑**与 2007 年中国固投增速见顶的历史经验直接相关。"
    "中国固投增速从 25% 跌到 20%（5 个百分点）用了 1 个季度，引发了 2008 年的需求坍塌。"
    "M7 Capex/Rev 从 30% 跌到 25%（5 个百分点）可能同样剧烈——"
    "**关键观察点是 M7 CFO 在季报电话会中的「措辞变化」**。"
    "如果从「加大投资」变为「审慎评估」或「优化效率」，意味着拐点临近。",
    indent=True)

H3(doc, "NDX vs NVDA 归一化对比")

add_chart(doc, '06_ndx_nvda.png',
          '图表 7：NDX vs NVDA 归一化对比（2022-11=100）——NVDA 涨幅 8 倍于 NDX',
          width_inches=6.5)

P(doc, "图表 7 揭示了 23-26 浪潮的一个关键特征：**NVDA 涨幅 800% 远超 NDX 整体 95%**——"
    "这是 2005-2007 没有的现象（彼时云铜 +1,122%，但有色指数整体 +1,857%，"
    "板块涨幅大于单股），这意味着 AI 浪潮的「龙头集中度」远高于 2005-2007。",
    indent=False)

P(doc, "这一集中度带来的影响是双面的：一方面，NVDA 的 P/S 估值（~20x）"
    "是整个 AI 板块的「锚」——如果 NVDA 估值崩溃，整个板块将连锁反应；"
    "另一方面，NVDA 2025-2026 业绩的高增长（EPS +147%）正在「消化」高估值，"
    "**S3 信号的触发需要看到「业绩增长放缓 + 估值高位」的双重确认**。",
    indent=False)


# ============================================================
# 第三部分
# ============================================================
H1(doc, "第三部分　三大量化顶部信号与操作框架")

P(doc,
    "基于第一部分的 6 维度复盘和第二部分的 10 维度对标，我们提出 3 个可量化、可监测、可证伪的顶部信号。"
    "每个信号对应一种「终结机制」，组合使用可在不同情境下提示风险。",
    indent=False)

P(doc, "**信号设计原则**：", indent=False)

for line in [
    "**可量化**：每个信号有明确的数值阈值（如 Capex/Rev 35%）",
    "**可监测**：每个信号有明确的观测事件和数据来源",
    "**可证伪**：每个信号有清晰的失效场景",
    "**领先性**：尽量选择领先指标，而非同步或滞后指标",
    "**互补性**：3 个信号对应 3 种不同的终结机制，互不重叠",
]:
    p = doc.add_paragraph(line, style='List Bullet')
    p.paragraph_format.left_indent = Pt(40)


H2(doc, "§3.1　S1：云厂商 Capex 增速二阶导转负")

H3(doc, "信号定义")

add_table(doc, [
    ["项目", "内容"],
    ["核心逻辑", "算力需求的唯一「真金」指标瓦解；证伪永续高增长"],
    ["历史对标", "2007.11 中国固定资产投资增速见顶 → 2008 铜真实需求断崖"],
    ["阈值", "M7 Capex 同比增速连续两季 <20% 或环比负增长"],
    ["当前状态", "**未触发**（M7 Capex/Rev 30%，增速仍高，距 35% 阈值 5pct）"],
    ["性质", "⏱ 领先指标（领先 1-2 季度）"],
])

H3(doc, "观测事件与数据来源")

P(doc, "**S1 信号的观测依赖两个核心事件**：", indent=False)

P(doc, "**(1) M7 季度财报中的 Capex 指引**。重点关注 Microsoft CFO Amy Hood、"
    "Google CFO Ruth Porat、Amazon CFO Brian Olsavsky、Meta CFO Susan Li 的表述变化。"
    "如果从「加大投资」转向「审慎评估」或「优化效率」，意味着拐点临近。", indent=True)

P(doc, "**(2) 台积电 CoWoS 先进封装订单变化**。CoWoS 是 AI GPU 的关键封装工艺，"
    "台积电的 CoWoS 订单数据是 M7 Capex 的「先行指标」——M7 下单后 1-2 个季度才体现在 CoWoS 产能上。"
    "如果 CoWoS 订单出现「环比下滑」，意味着 M7 Capex 增速已实质放缓。", indent=True)

P(doc, "**数据来源**：SEC 10-Q/10-K 公开披露；台积电月度营收公告；", indent=False)
P(doc, "**更新频率**：季度（财报季）+ 月度（台积电营收）", indent=False)


H2(doc, "§3.2　S2：标志性 AI 独角兽巨型 IPO 首日即巅峰")

H3(doc, "信号定义")

add_table(doc, [
    ["项目", "内容"],
    ["核心逻辑", "市场流动性已无法支撑「梦想变现」；信心与比价崩塌"],
    ["历史对标", "2007.11.05 中石油上市高开低走 → 上证指数大顶"],
    ["阈值", "IPO 首日高开 >50% 后收长阴线（收盘价 < 开盘价）"],
    ["当前状态", "**未触发**（OpenAI 未上市）"],
    ["性质", "🔴 同步确认信号"],
])

H3(doc, "OpenAI 拟 IPO 估值变化")

add_table(doc, [
    ["时间", "估值", "事件"],
    ["2025-10", "$500B（要约收购）", "tender offer"],
    ["2026-03", "$850B（内部估值）", "估值快速攀升"],
    ["2026-06", "拟 IPO $100B", "预计 2026Q4-2027Q1"],
])

H3(doc, "S2 触发后的市场反应预判")

P(doc, "OpenAI 上市后 30 日的表现是**决定性观察点**。", indent=False)

P(doc, "**情景 A：上市首日高开 +20% 以内，收盘稳定**。这是健康的市场表现，"
    "说明市场对 AI 估值仍有支撑。S2 信号**未触发**。", indent=True)

P(doc, "**情景 B：上市首日高开 +30-50%，收盘稳定**。这是偏热但不极端的表现，"
    "需要继续观察 30 日内的「二次发行」情况。S2 信号**预警但未触发**。", indent=True)

P(doc, "**情景 C：上市首日高开 +50% 以上，收盘低于开盘价（长阴线）**。"
    "这是 S2 信号**触发**的场景。意味着市场已无法承接「梦想变现」的体量，"
    "原始股东和风投的抛售压力将引发连锁反应——所有 AI 独角兽的估值需重新定价。", indent=True)

P(doc, "**与中石油 IPO 的对比**：2007-11-05 中石油 A 股首日 +191%，随后 -22%（30 日）、-65%（1 年）。"
    "如果 OpenAI 重现这一模式，AI 板块的顶部将得到「同步确认」。", indent=False)


H2(doc, "§3.3　S3：龙头估值框架崩溃 + 应用端收入证伪")

H3(doc, "信号定义")

add_table(doc, [
    ["项目", "内容"],
    ["核心逻辑", "估值从「梦想定价」瞬间切换为「残值定价」；应用端不产生真金白银增量"],
    ["历史对标", "2007.10 中铝 PE 65x，靠「注矿」故事；2008 年业绩暴跌，股价跌超 90%"],
    ["阈值", "① 英伟达跌破 200 日均线 ② 微软 Copilot 付费渗透率停滞"],
    ["当前状态", "**未触发**（NVDA P/S ~20，Copilot 渗透率 15% 接近阈值）"],
    ["性质", "🔵 滞后/确认信号"],
])

H3(doc, "S3 双重阈值的逻辑")

P(doc, "**S3 信号需要看到「估值崩溃」+「收入证伪」的双重确认**，理由如下：", indent=False)

P(doc, "**单独看 NVDA 估值**可能产生误判。NVDA 当前 P/S ~20x，处于历史 60% 分位。"
    "如果仅看估值，NVDA 似乎仍有上涨空间。但 P/S 的「合理」是建立在「营收持续高增长」前提上的。"
    "**如果营收增速降至个位数，20x P/S 就会显得过高**——估值切换可能在一个月内完成。", indent=True)

P(doc, "**单独看 Copilot 渗透率**也不足以触发 S3。Copilot 当前 15%，距离 25-30% 的「成熟期」"
    "还有翻倍空间。**真正的证伪是「渗透率停滞」而非「渗透率低」**——"
    "如果连续 2 个季度停滞在 15%，意味着 AI 应用端的「故事」已被证伪。", indent=True)

P(doc, "**双重确认的机制**：当 NVDA 估值与 Copilot 渗透率同时出现问题时，"
    "市场会意识到「算力侧和应用侧」同时见顶——这是 AI 浪潮的系统性顶部信号。", indent=False)


H2(doc, "§3.4　综合评分与操作框架")

add_chart(doc, '08_signal_dashboard.png',
          '图表 8：三大量化顶部信号当前状态仪表盘（2026-06）',
          width_inches=6.5)

P(doc, "图表 8 以仪表盘形式呈现了 3 个信号的当前状态。**综合评分 0/3 ~ 0.5/3**——"
    "S1 接近触发但未触发（Capex/Rev 30% 距阈值 5pct），S2/S3 距离触发仍有空间。",
    indent=False)

P(doc, "**操作框架与评分对应**：", indent=False)

add_table(doc, [
    ["综合评分", "状态", "操作建议"],
    ["0/3 ~ 0.5/3", "**安全**（当前）", "维持 AI 仓位 60-70%，跟踪 S1"],
    ["1/3", "**预警**", "减仓 AI 至 30-40%，提现金"],
    ["2/3", "**危险**", "AI 降至 10-20%，转配债券/黄金"],
    ["3/3", "**触发**", "清仓风险资产，等右侧"],
])

P(doc, "**关键观察点**：OpenAI IPO 后 30 日 = 决定性窗口。", indent=False)


# ============================================================
# 附录
# ============================================================
doc.add_page_break()

H1(doc, "附录 A　三国货币政策指标不对称的方法论说明")

P(doc, "本附录系统说明三国货币政策指标不对称的根源、修正方法。", indent=False)

add_table(doc, [
    ["国家", "央行目标", "关键指标", "原因"],
    ["🇺🇸 美国", "充分就业 + 物价稳定（双目标制）", "联邦基金利率", "利率走廊机制（IORB + ON RRP）"],
    ["🇨🇳 中国", "经济增长 + 货币供应（多目标制）", "M2 + 1Y 贷款/存款基准", "数量型工具 + 管制利率"],
    ["🇯🇵 日本", "通胀 + 利率曲线（量宽后转价格型）", "无担保隔夜利率", "量化宽松后转向价格型"],
])

H2(doc, "A.1　4 项方法论修正")

P(doc, "**修正 1：必须用「同期限市场利率」做「短长利差」对比**", indent=False)

P(doc, "原版报告曾出现「中国 1Y 贷款基准 7.47% vs 10Y 国债 4.5%，短长利差 -3.0%」"
    "的错误对比——贷款基准是央行管制利率（非市场化），国债是市场化利率，两者本质不同，"
    "不能直接比较「短长利差」。正确做法是用同期限的市场利率做对比："
    "1Y 国债 2.79% vs 10Y 国债 4.5% → 市场短长利差 +1.71%（正常）。", indent=True)

P(doc, "**修正 2：必须区分「政策利率」和「市场利率」**", indent=False)

P(doc, "政策利率是央行直接规定的利率（如中国 1Y 贷款基准 7.47% / 1Y 存款基准 4.14%），"
    "市场利率是金融市场交易形成的利率（如中国 1Y 国债 2.79% / 10Y 国债 4.5%）。"
    "两者的关系是「政策利率影响市场利率，但市场利率反映真实金融条件」。", indent=True)

P(doc, "**修正 3：中国「管制利率倒挂」是 2007 年特有现象**", indent=False)

P(doc, "1Y 存款基准 4.14% > 1Y 国债 2.79%——在正常利率体系中，国债作为无风险利率应该是市场利率的下限，"
    "但中国央行通过存款管制将存款利率人为抬高在国债之上。"
    "这一倒挂在 2015 年存款利率上限取消后逐步消失。", indent=True)

P(doc, "**修正 4：贷存利差 3.33% = 银行业「制度性红包」**", indent=False)

P(doc, "1Y 贷款基准 7.47% - 1Y 存款基准 4.14% = 3.33%。"
    "在 2019 年 LPR 改革前，这是央行给中国银行业的制度性保护——"
    "3.33% 的安全垫是银行利润的核心来源。"
    "2019 年 LPR 改革后，贷存利差收窄至 1.5%。", indent=True)

quote_block(doc,
    "**结论**：中国货币政策是「数量型 + 管制利率」二元体系。"
    "需用「政策利率 + 市场利率 + 贷存利差 + 外汇占款」四维度才能完整刻画。")


# 附录 B
H1(doc, "附录 B　05-07 重大事件时间轴")

P(doc, "本附录以时间顺序列出 2005-2007 期间对市场产生重大影响的 16 个事件，"
    "供读者对照图表 3、图表 4 的事件标注查阅。", indent=False)

add_table(doc, [
    ["日期", "事件", "类型", "影响维度"],
    ["2005.06", "股改启动", "政策", "资本市场"],
    ["2005.07", "人民币汇改", "汇率", "汇率"],
    ["2005.11", "神华 H 股 IPO", "IPO", "IPO"],
    ["2006.04", "中铝 A 股 IPO", "IPO", "IPO"],
    ["2006.05", "LME 铜 $8,800 历史新高", "商品", "行业"],
    ["2006.06", "美联储加息见顶 5.25%", "货币", "货币政策"],
    ["2007.04.30", "中铝吸收合并山铝/兰铝", "公司", "公司"],
    ["2007.05.30", "印花税 1‰→3‰", "政策", "政策"],
    ["2007.07.12", "西部矿业上市", "IPO", "IPO"],
    ["2007.09.25", "建设银行 A 股 IPO", "IPO", "IPO"],
    ["2007.10.09", "中国神华 A 股 IPO", "IPO", "IPO"],
    ["2007.10.16", "上证 6124 见顶", "市场", "资本市场"],
    ["2007.11.05", "中石油 A 股 IPO", "IPO", "IPO"],
    ["2008.07.03", "LME 铜 $8,985 见顶", "商品", "行业"],
    ["2008.10.28", "上证 1664 金融危机底", "市场", "资本市场"],
    ["2008.12", "雷曼效应", "宏观", "危机"],
])


# 附录 C
H1(doc, "附录 C　数据来源")

P(doc, "本报告数据来源说明：", indent=False)

data_sources = [
    "**FRED** (Federal Reserve Economic Data)：美联储官方经济数据",
    "**PBOC** (People's Bank of China)：中国人民银行",
    "**BOJ** (Bank of Japan)：日本银行",
    "**IMF** (International Monetary Fund)：国际货币基金组织",
    "**World Bank**：世界银行",
    "**SEC EDGAR**：美国证券交易委员会公开数据",
    "**LME** (London Metal Exchange)：伦敦金属交易所",
    "**ICSG** (International Copper Study Group)：国际铜业研究组",
    "**海关总署**：中国海关",
    "**Bloomberg**：金融数据终端",
]
for s in data_sources:
    doc.add_paragraph(s, style='List Bullet')

P(doc, "**核心数据文件清单**（TopicResearch/data/raw/）：", indent=False)

P(doc, "本报告使用了 32 个 JSON 数据文件，包括 LME 铜铝日级数据、上证综指日级数据、"
    "FRED 联邦基金利率/10Y 国债/30Y 房贷/美元指数/M2/企业债利差、PBOC 中国 M2/外储/利率/贸易顺差、"
    "BOJ 日本利率/JGB/汇率/外储、AI 板块 M7 营收/Capex/NDX/Copilot、05-07 IPO 数据、"
    "金属涨幅数据、重大事件时间轴、IMF/卖方预期、中国大盘数据等。", indent=True)


# 末尾
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("—— 报告结束 ——")
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = RGBColor(148, 163, 184)


# 保存
doc.save(OUTPUT)
print(f"✅ 已生成: {OUTPUT}")
print(f"   文件大小: {os.path.getsize(OUTPUT)/1024:.1f} KB")
print(f"   段落数: {len(doc.paragraphs)}")
print(f"   表格数: {len(doc.tables)}")
