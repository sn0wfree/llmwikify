"""
更新 Word 文档宏观章节：添加图表 + 精简文字
"""
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCX_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.5.docx'
CHART_DIR = '/home/ll/llmwikify/TopicResearch/report/charts'
OUTPUT_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'

# 图表文件名映射
CHARTS = {
    'gdp_nominal': 'chart_gdp_nominal.png',
    'gdp_real': 'chart_gdp_real.png',
    'cpi': 'chart_cpi.png',
    'monetary': 'chart_monetary.png',
    'term_spread': 'chart_term_spread.png',
    'fx_trade': 'chart_fx_trade.png',
    'forex_reserve': 'chart_forex_reserve.png',
    'impossible_trinity': 'chart_impossible_trinity.png',
}

# 精简后的文字内容
TEXT_CONTENT = {
    'gdp_nominal_chart_title': '图 1.1：名义 GDP 增速对比（2000-2026）',
    'gdp_nominal_text': (
        '2005-2007 年中国名义 GDP 增速从 15.6% 升至 23.2%，显著高于美国（6.7→4.8%）和日本（0.5→0.7%）。'
        '中国名义 GDP 增速显著高于实际增速（14%），差值源于通胀及资产价格重估。'
        '数据来源：ifind EDB（名义 GDP 本币口径）。'
    ),
    'gdp_real_chart_title': '图 1.2：实际 GDP 增速对比（2000-2026）',
    'gdp_real_text': (
        '剔除价格因素后，中国实际 GDP 增速 11-14%，美国 2-3%，日本 1-2%。'
        '三国经济周期并不同步但结构互补：中国投资驱动、美国消费驱动、日本出口+套利。'
        '2008 年金融危机后，美国实际 GDP 跌至 -2.6%，复苏缓慢。'
        '数据来源：ifind EDB（实际 GDP 不变价）。'
    ),
    'cpi_chart_title': '图 1.3：CPI 通胀对比（2000-2026）',
    'cpi_text': (
        '2005-2007 年中国 CPI 从 1.8% 升至 4.8%，通胀温和可控；美国 CPI 维持 2-3%；日本长期通缩（CPI ≈0）。'
        '2008 年危机后全球通缩，2022 年美欧通胀飙升至 8%+，日本 2022 年才摆脱通缩。'
        '数据来源：ifind EDB。'
    ),
    'monetary_chart_title': '图 1.4：中国货币政策——表面紧缩与实际宽松的背离',
    'monetary_text': (
        '2006-2007 年 PBOC 加息 6 次（5.58→7.47%）+ 提 RRR 10 次（7.5→14.5%），但 M2 维持 17-19%。'
        '原因：贸易顺差 1020→2639 亿美元 → 外汇占款被动投放基础货币 → 紧缩被对冲。'
        '2019 年 LPR 改革后，货币政策从数量型转向价格型。'
        '数据来源：ifind EDB（利率、RRR、M2、贸易顺差）。'
    ),
    'term_spread_chart_title': '图 1.5：短端/长端市场利率与货币政策类型',
    'term_spread_text': (
        '美国——价格型：以联邦基金利率为中介目标，通过公开市场操作调节短端利率（泰勒规则）。'
        '中国——数量型（2019 前）→ 价格型（2019 后）：以 M2 增速为中介目标，2019 年 LPR 改革后转向利率走廊。'
        '日本——价格型（YCC）：以 10Y 国债收益率为锚，短端维持负利率。'
        '数据来源：ifind EDB（银行间回购利率、国债收益率）。'
    ),
    'fx_trade_chart_title': '图 1.6：汇率与贸易——全球失衡与资金流向',
    'fx_trade_text': (
        '2005.7.21 汇改：人民币一次性升值 2.1%（8.28→8.11），此后渐进升值至 2008 年 6.83。'
        'DXY 2002-2007 从 120 跌至 80（贬值 33%），大宗商品以美元计价上涨。'
        '中国贸易顺差 1020→2639 亿美元（2005-2007），美国逆差维持 600 亿$/月。'
        '数据来源：ifind EDB（汇率、贸易差额）。'
    ),
    'forex_reserve_chart_title': '图 1.7：外汇储备/GDP——从积累到消耗',
    'forex_reserve_text': (
        '中国外储/GDP 从 2000 年 15% 升至 2014 年峰值 49%（3.99 万亿美元），2015-2016 年资本外流消耗近 1 万亿美元。'
        '美国外储/GDP 仅 0.3%（美元储备货币地位），日本约 20%。'
        '当前中国外储稳定在 3.2 万亿，GDP 增长稀释比率至 17%。'
        '数据来源：ifind EDB（外储、GDP）。'
    ),
    'impossible_trinity_title': '图 1.8：蒙代尔不可能三角与中国政策选择',
    'impossible_trinity_text': (
        '蒙代尔不可能三角：固定汇率、资本自由流动、独立货币政策三者不可兼得。'
        '中国 2005-2007 选择 B 组合（固定汇率+名义管制），代价是冲销成本 1-2% GDP。'
        '当前向 C 组合移动（更灵活汇率+更独立货币政策），外储从 4 万亿降至 3.2 万亿。'
        '美国处于 C 组合（浮动汇率+资本自由+独立货币政策），得益于美元储备货币地位。'
    ),
}

def add_chart(doc, heading_text, chart_key, caption, text, after_text=None):
    """在指定标题后添加图表和文字"""
    # 找到目标标题
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if heading_text in p.text and p.style.name.startswith('Heading'):
            target_idx = i
            break
    
    if target_idx is None:
        print(f'  ⚠️ 未找到标题: {heading_text}')
        return False
    
    # 找到下一个标题的位置
    next_idx = len(doc.paragraphs)
    for i in range(target_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            next_idx = i
            break
    
    # 在下一个标题前插入内容
    # 注意：python-docx 不支持直接在指定位置插入，需要在末尾添加后移动
    # 简化方案：在文档末尾添加，由用户手动调整位置
    
    # 添加图表标题
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加图表
    chart_path = os.path.join(CHART_DIR, CHARTS[chart_key])
    if os.path.exists(chart_path):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加文字说明
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    
    print(f'  ✅ {caption}')
    return True

def main():
    print('Loading document...')
    doc = docx.Document(DOCX_PATH)
    print(f'  段落数: {len(doc.paragraphs)}')
    
    # 由于 python-docx 无法在指定位置插入，改用追加方式
    # 在文档末尾添加宏观背景章节
    
    print('\nAdding macro background section...')
    
    # 添加分页符
    doc.add_page_break()
    
    # 添加章节标题
    p = doc.add_heading('附录 D：宏观经济背景图表（Slides 1.1.1-1.2.5）', level=1)
    
    # 图表 1：名义 GDP
    p = doc.add_heading('D.1 名义 GDP 增速对比', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.1：名义 GDP 增速对比（2000-2026）')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['gdp_nominal'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['gdp_nominal_text'])
    run.font.size = Pt(9)
    
    # 图表 2：实际 GDP
    p = doc.add_heading('D.2 实际 GDP 增速对比', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.2：实际 GDP 增速对比（2000-2026）')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['gdp_real'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['gdp_real_text'])
    run.font.size = Pt(9)
    
    # 图表 3：CPI
    p = doc.add_heading('D.3 CPI 通胀对比', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.3：CPI 通胀对比（2000-2026）')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['cpi'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['cpi_text'])
    run.font.size = Pt(9)
    
    # 图表 4：货币政策
    p = doc.add_heading('D.4 货币政策背离', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.4：中国货币政策——表面紧缩与实际宽松的背离')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['monetary'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['monetary_text'])
    run.font.size = Pt(9)
    
    # 图表 5：市场利率
    p = doc.add_heading('D.5 市场利率与货币政策类型', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.5：短端/长端市场利率与货币政策类型')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['term_spread'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['term_spread_text'])
    run.font.size = Pt(9)
    
    # 图表 6：汇率与贸易
    p = doc.add_heading('D.6 汇率与贸易', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.6：汇率与贸易——全球失衡与资金流向')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['fx_trade'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['fx_trade_text'])
    run.font.size = Pt(9)
    
    # 图表 7：外汇储备
    p = doc.add_heading('D.7 外汇储备/GDP', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.7：外汇储备/GDP——从积累到消耗')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['forex_reserve'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(6))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['forex_reserve_text'])
    run.font.size = Pt(9)
    
    # 图表 8：不可能三角
    p = doc.add_heading('D.8 蒙代尔不可能三角', level=2)
    p = doc.add_paragraph()
    run = p.add_run('图 1.8：蒙代尔不可能三角与中国政策选择')
    run.font.size = Pt(10)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_path = os.path.join(CHART_DIR, CHARTS['impossible_trinity'])
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(chart_path, width=Inches(5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    run = p.add_run(TEXT_CONTENT['impossible_trinity_text'])
    run.font.size = Pt(9)
    
    # 保存
    print(f'\nSaving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'✅ Done! Saved to {OUTPUT_PATH}')

if __name__ == '__main__':
    main()
