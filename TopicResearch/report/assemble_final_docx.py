"""
组装最终文档 v2.8：
1. 基于 v2.5（保留所有37个表格）
2. 在 1.3.4 后插入 AI IPO 抽血分析
3. 在对应章节插入宏观图表（正文位置）
4. 保存为 v2.8
"""
import os
import copy
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.expanduser("~/llmwikify/TopicResearch/report")
CHARTS_DIR = os.path.join(BASE_DIR, "charts")
V25_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.5.docx")
OUT_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.8.docx")

# Word XML namespace
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# 图表-章节对应关系：(图表文件名, 章节标题关键词, 图表标题)
CHART_MAPPINGS = [
    ("chart_gdp_nominal.png", "1.1.1 经济增长", "图 1.1：名义 GDP 增速对比（2000-2026）"),
    ("chart_gdp_real.png", "1.1.1 经济增长", "图 1.2：实际 GDP 增速对比（2000-2026）"),
    ("chart_cpi.png", "1.1.2 货币政策", "图 1.3：CPI 通胀对比（2000-2026）"),
    ("chart_monetary.png", "1.1.2 货币政策", "图 1.4：中国货币政策背离"),
    ("chart_term_spread.png", "1.1.2 货币政策", "图 1.5：市场利率与货币政策类型"),
    ("chart_fx_trade.png", "1.1.3 汇率与贸易", "图 1.6：汇率与贸易"),
    ("chart_forex_reserve.png", "1.1.4 外汇储备", "图 1.7：外汇储备/GDP"),
    ("chart_impossible_trinity.png", "1.4.2 三角闭环", "图 1.8：蒙代尔不可能三角"),
]


def get_element_text(elem):
    """获取元素的文本内容"""
    text = ''
    if elem.text:
        text += elem.text
    for child in elem:
        if child.text:
            text += child.text
        if child.tail:
            text += child.tail
    return text


def find_element_index(body, keyword):
    """在 body 中找到包含关键词的段落索引"""
    for i, elem in enumerate(body):
        if elem.tag == f'{{{nsmap["w"]}}}p':
            text = get_element_text(elem)
            if keyword in text:
                return i
    return -1


def create_paragraph_element(text, bold=False, size=9, alignment=None):
    """创建段落 XML 元素"""
    p = etree.Element(f'{{{nsmap["w"]}}}p')
    
    # 段落属性
    pPr = etree.SubElement(p, f'{{{nsmap["w"]}}}pPr')
    
    if alignment:
        jc = etree.SubElement(pPr, f'{{{nsmap["w"]}}}jc')
        jc.set(f'{{{nsmap["w"]}}}val', alignment)
    
    # Run
    r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
    
    # Run 属性
    rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
    rFonts = etree.SubElement(rPr, f'{{{nsmap["w"]}}}rFonts')
    rFonts.set(f'{{{nsmap["w"]}}}ascii', '微软雅黑')
    rFonts.set(f'{{{nsmap["w"]}}}hAnsi', '微软雅黑')
    
    sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
    sz.set(f'{{{nsmap["w"]}}}val', str(size * 2))
    
    if bold:
        b = etree.SubElement(rPr, f'{{{nsmap["w"]}}}b')
    
    # 文本
    t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return p


def create_table_element(table_data):
    """创建表格 XML 元素"""
    rows = len(table_data)
    cols = len(table_data[0]) if table_data else 0
    
    tbl = etree.Element(f'{{{nsmap["w"]}}}tbl')
    
    tblPr = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tblPr')
    tblStyle = etree.SubElement(tblPr, f'{{{nsmap["w"]}}}tblStyle')
    tblStyle.set(f'{{{nsmap["w"]}}}val', 'TableGrid')
    tblW = etree.SubElement(tblPr, f'{{{nsmap["w"]}}}tblW')
    tblW.set(f'{{{nsmap["w"]}}}w', '0')
    tblW.set(f'{{{nsmap["w"]}}}type', 'auto')
    
    tblGrid = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tblGrid')
    for _ in range(cols):
        etree.SubElement(tblGrid, f'{{{nsmap["w"]}}}gridCol')
    
    for row_idx, row_data in enumerate(table_data):
        tr = etree.SubElement(tbl, f'{{{nsmap["w"]}}}tr')
        for col_idx, cell_text in enumerate(row_data):
            tc = etree.SubElement(tr, f'{{{nsmap["w"]}}}tc')
            p = etree.SubElement(tc, f'{{{nsmap["w"]}}}p')
            r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
            rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
            sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
            sz.set(f'{{{nsmap["w"]}}}val', '16')
            if row_idx == 0:
                b = etree.SubElement(rPr, f'{{{nsmap["w"]}}}b')
            t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
            t.text = str(cell_text)
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return tbl


def create_ai_ipo_elements():
    """创建 AI IPO 分析的 XML 元素列表"""
    elements = []
    
    elements.append(create_paragraph_element(
        "1.3.4-AI IPO抽血分析（本轮AI巨型IPO）",
        size=10, bold=True))
    
    elements.append(create_paragraph_element(
        '2026年6月，SpaceX在纳斯达克成功上市（股票代码SPCX），募资750亿美元，创下全球资本市场史上最大IPO纪录。'
        '紧随其后，OpenAI计划于2026Q4-2027Q1进行IPO，规模预计1000亿美元。'
        '本轮AI巨兽集中上市的"抽血效应"值得高度关注。',
        size=9))
    
    elements.append(create_paragraph_element(
        "表 1.3-1：本轮AI巨型IPO全景",
        size=9, bold=True, alignment='center'))
    
    elements.append(create_table_element([
        ['公司', '状态', '募资规模', '估值', '抽血占比'],
        ['SpaceX', '✅ 已上市 (2026.6.12)', '$750亿', '$1.77万亿', '0.17%'],
        ['OpenAI', '拟IPO (2026Q4-2027Q1)', '$1000亿', '$8500亿', '0.22%'],
        ['Anthropic', '拟IPO (2027+)', '$500亿', '$2000亿', '0.11%'],
        ['合计', '-', '$2250亿', '-', '0.50%'],
    ]))
    
    elements.append(create_paragraph_element(
        "表 1.3-2：SpaceX vs 中石油 IPO对比",
        size=9, bold=True, alignment='center'))
    
    elements.append(create_table_element([
        ['维度', '2007 中石油', '2026 SpaceX', '差异'],
        ['募资规模', '¥668亿 ($97亿)', '$750亿', 'SpaceX是中石油7.7倍'],
        ['上市估值', '$1.08万亿', '$1.77万亿', 'SpaceX高64%'],
        ['首日涨幅', '+191% (开盘)', '+29% (开盘)', '中石油更疯狂'],
        ['冻结资金', '¥3.3万亿 (40%)', '$2500亿认购 (4倍)', '中石油抽血更猛'],
        ['散户参与', '>60%', '30%', '散户占比下降'],
        ['盈利状态', '盈利', '亏损 ($49.4亿)', 'SpaceX尚未盈利'],
        ['后续表现', '一年内破发', '待观察', '-'],
    ]))
    
    elements.append(create_paragraph_element(
        "AI IPO抽血机制分析：",
        size=9, bold=True))
    
    analyses = [
        ('（1）抽血规模对比：', 'SpaceX募资750亿美元，占标普500流通市值约0.17%；OpenAI预计募资1000亿美元，占纳斯达克100流通市值约0.4%。合计潜在抽血2250亿美元，占标普500流通市值约0.50%。对比2007年中石油冻结3.3万亿元（占A股流通市值40%），本轮抽血规模远小于2007年。'),
        ('（2）存量占用更严重：', '本轮AI牛市中，NVIDIA+Microsoft已占纳斯达克100约24%权重。SpaceX上市后迅速纳入纳斯达克100，进一步加剧头部集中。当AI巨兽占据指数主要权重时，新IPO的边际抽血效应可能被放大。'),
        ('（3）流动性环境差异：', '2007年M2增速17%+外汇占款被动投放，流动性极度宽松；2026年M2增速8%+美联储维持高利率，流动性相对紧张。这意味着本轮IPO对市场流动性的冲击可能更敏感。'),
        ('（4）散户参与度下降：', 'SpaceX散户份额30%（远高于近年平均水平），但中石油时代散户占比超60%。机构主导的IPO通常波动更小，但一旦下跌，机构止损盘可能更猛烈。'),
    ]
    
    for title, text in analyses:
        elements.append(create_paragraph_element(
            title + text,
            size=9))
    
    elements.append(create_paragraph_element(
        "关键判断：",
        size=9, bold=True))
    
    elements.append(create_paragraph_element(
        '本轮AI IPO抽血规模虽小于2007年，但需高度关注两个风险点：'
        '（1）多只AI巨兽集中上市的时间窗口重叠（SpaceX已上市+OpenAI即将IPO）；'
        '（2）存量占用严重（NVIDIA+Microsoft+SpaceX已占纳斯达克100约26%）。'
        '核心监测指标：OpenAI IPO后30日表现（S2信号触发点）。',
        size=9))
    
    return elements


def main():
    print("Loading v2.5 document...")
    doc = Document(V25_PATH)
    print(f"  段落数: {len(doc.paragraphs)}")
    print(f"  表格数: {len(doc.tables)}")
    
    body = doc.element.body
    
    # 找到 1.3.4 章节位置
    s134_idx = find_element_index(body, "1.3.4 巨型IPO现象")
    if s134_idx == -1:
        print("ERROR: Cannot find section 1.3.4")
        return
    
    print(f"  1.3.4 section at element: {s134_idx}")
    
    # 找到 1.4 章节位置
    s14_idx = find_element_index(body, "1.4 超级周期的汇聚与终结机制")
    if s14_idx == -1:
        print("ERROR: Cannot find section 1.4")
        return
    
    print(f"  1.4 section at element: {s14_idx}")
    
    # 找到各章节位置
    section_positions = {}
    for keyword in ["1.1.1 经济增长", "1.1.2 货币政策", "1.1.3 汇率与贸易", 
                     "1.1.4 外汇储备", "1.4.2 三角闭环"]:
        idx = find_element_index(body, keyword)
        if idx != -1:
            section_positions[keyword] = idx
            print(f"  {keyword} at element: {idx}")
    
    # 插入 AI IPO 分析元素
    print("\nInserting AI IPO analysis...")
    ai_ipo_elements = create_ai_ipo_elements()
    
    content_start = s134_idx + 1
    for i, elem in enumerate(ai_ipo_elements):
        body.insert(content_start + i, elem)
    
    print(f"  Inserted {len(ai_ipo_elements)} elements")
    
    # 记录图表插入位置（使用标记段落）
    print("\nAdding chart placeholders...")
    
    # 重新查找章节位置
    section_positions_updated = {}
    for keyword in ["1.1.1 经济增长", "1.1.2 货币政策", "1.1.3 汇率与贸易", 
                     "1.1.4 外汇储备", "1.4.2 三角闭环"]:
        idx = find_element_index(body, keyword)
        if idx != -1:
            section_positions_updated[keyword] = idx
    
    # 按章节分组
    charts_by_section = {}
    for chart_file, section_key, caption in CHART_MAPPINGS:
        if section_key not in charts_by_section:
            charts_by_section[section_key] = []
        charts_by_section[section_key].append((chart_file, caption))
    
    # 从后往前插入标记
    for section_key, charts in charts_by_section.items():
        if section_key in section_positions_updated:
            section_idx = section_positions_updated[section_key]
            next_sections = {
                "1.1.1 经济增长": "1.1.2 货币政策",
                "1.1.2 货币政策": "1.1.3 汇率与贸易",
                "1.1.3 汇率与贸易": "1.1.4 外汇储备",
                "1.1.4 外汇储备": "1.1.5 财政政策",
                "1.4.2 三角闭环": "1.4.3 三大见顶特征",
            }
            next_section = next_sections.get(section_key)
            if next_section:
                end_idx = find_element_index(body, next_section)
                if end_idx == -1:
                    end_idx = len(body) - 1
            else:
                end_idx = len(body) - 1
            
            # 在章节末尾插入标记
            for chart_file, caption in reversed(charts):
                chart_path = os.path.join(CHARTS_DIR, chart_file)
                if os.path.exists(chart_path):
                    # 插入标题
                    title_elem = create_paragraph_element(caption, bold=True, size=9, alignment='center')
                    body.insert(end_idx, title_elem)
                    
                    # 插入图片占位符（带特殊标记）
                    placeholder = create_paragraph_element(
                        f"CHART:{chart_file}",
                        size=9, alignment='center')
                    body.insert(end_idx + 1, placeholder)
                    
                    print(f"  ✅ {caption} → {section_key}")
    
    # 保存文档
    print(f"\nSaving to {OUT_PATH}...")
    doc.save(OUT_PATH)
    
    # 使用 python-docx 重新打开并替换占位符为实际图片
    print("\nReplacing placeholders with actual images...")
    doc2 = Document(OUT_PATH)
    
    # 遍历所有段落，找到占位符并替换
    for i, p in enumerate(doc2.paragraphs):
        if p.text.startswith("CHART:"):
            chart_file = p.text.replace("CHART:", "")
            chart_path = os.path.join(CHARTS_DIR, chart_file)
            
            if os.path.exists(chart_path):
                # 清除占位符文本
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加图片
                run = p.add_run()
                run.add_picture(chart_path, width=Inches(5.5))
                
                print(f"  ✅ Inserted: {chart_file}")
    
    # 保存最终文档
    doc2.save(OUT_PATH)
    
    # 验证
    verify_doc = Document(OUT_PATH)
    print(f"\nVerification:")
    print(f"  段落数: {len(verify_doc.paragraphs)}")
    print(f"  表格数: {len(verify_doc.tables)}")
    
    # 检查图片
    image_count = sum(1 for rel in verify_doc.part.rels.values() if 'image' in rel.reltype)
    print(f"  图片数: {image_count}")
    
    print(f"  ✅ Done!")


if __name__ == "__main__":
    main()
