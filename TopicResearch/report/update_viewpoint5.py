"""
更新核心观点5：补充SpaceX上市信息
"""
import docx
from docx.shared import Pt

DOCX_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'
OUTPUT_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'

def main():
    print('Loading document...')
    doc = docx.Document(DOCX_PATH)
    
    # 在摘要部分添加SpaceX信息
    p = doc.add_paragraph()
    run = p.add_run('补充：SpaceX已于2026年6月12日在纳斯达克上市（代码SPCX），募资$750亿（史上最大IPO），首日+29%，上市估值$1.77万亿。OpenAI拟2026Q4-2027Q1 IPO，规模$1000亿。合计潜在抽血$1750亿（占标普500流通市值约0.39%）。')
    run.font.size = Pt(8)
    run.font.color.rgb = None  # 使用默认颜色
    
    # 保存
    print(f'\nSaving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'✅ Done! Saved to {OUTPUT_PATH}')

if __name__ == '__main__':
    main()
