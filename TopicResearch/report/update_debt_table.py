"""
更新文档表格9（三国债务结构）添加2022-2026年数据
"""
import os
from lxml import etree
from docx import Document

BASE_DIR = os.path.expanduser("~/llmwikify/TopicResearch/report")
V29_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.9.docx")
OUT_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.10.docx")

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def get_element_text(elem):
    """获取元素的文本内容"""
    text = ''
    if elem.text:
        text += elem.text
    for child in elem:
        if child.text:
            text += child.text
        if child.tail:
            text += child.tail
    return text


def find_table_index(body, keyword):
    """找到包含关键词的表格索引"""
    tables = body.findall(f'.//{{{nsmap["w"]}}}tbl')
    for i, tbl in enumerate(tables):
        text = ''
        for p in tbl.findall(f'.//{{{nsmap["w"]}}}p'):
            text += get_element_text(p)
        if keyword in text:
            return i
    return -1


def update_table_cell(tbl, row_idx, col_idx, new_value):
    """更新表格单元格"""
    rows = tbl.findall(f'{{{nsmap["w"]}}}tr')
    if row_idx < len(rows):
        row = rows[row_idx]
        cells = row.findall(f'{{{nsmap["w"]}}}tc')
        if col_idx < len(cells):
            cell = cells[col_idx]
            # 清除原有内容
            for p in cell.findall(f'{{{nsmap["w"]}}}p'):
                cell.remove(p)
            # 添加新内容
            p = etree.SubElement(cell, f'{{{nsmap["w"]}}}p')
            r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
            rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
            sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
            sz.set(f'{{{nsmap["w"]}}}val', '18')
            t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
            t.text = new_value
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            return True
    return False


def add_table_column(tbl, col_idx, values):
    """在表格中插入新列"""
    rows = tbl.findall(f'{{{nsmap["w"]}}}tr')
    for i, row in enumerate(rows):
        cells = row.findall(f'{{{nsmap["w"]}}}tc')
        # 创建新单元格
        new_tc = etree.Element(f'{{{nsmap["w"]}}}tc')
        p = etree.SubElement(new_tc, f'{{{nsmap["w"]}}}p')
        r = etree.SubElement(p, f'{{{nsmap["w"]}}}r')
        rPr = etree.SubElement(r, f'{{{nsmap["w"]}}}rPr')
        sz = etree.SubElement(rPr, f'{{{nsmap["w"]}}}sz')
        sz.set(f'{{{nsmap["w"]}}}val', '18')
        if i == 0:  # 表头加粗
            b = etree.SubElement(rPr, f'{{{nsmap["w"]}}}b')
        t = etree.SubElement(r, f'{{{nsmap["w"]}}}t')
        t.text = values[i] if i < len(values) else ''
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        
        # 插入到指定位置
        if col_idx < len(cells):
            cells[col_idx].addprevious(new_tc)
        else:
            row.append(new_tc)


def main():
    print("Loading v2.9 document...")
    doc = Document(V29_PATH)
    print(f"  段落数: {len(doc.paragraphs)}")
    print(f"  表格数: {len(doc.tables)}")
    
    body = doc.element.body
    
    # 找到表格9（三国债务结构）
    tables = body.findall(f'.//{{{nsmap["w"]}}}tbl')
    table_9 = None
    for i, tbl in enumerate(tables):
        text = ''
        for p in tbl.findall(f'.//{{{nsmap["w"]}}}p'):
            text += get_element_text(p)
        if '美国政府债务/GDP' in text:
            table_9 = tbl
            print(f"  找到表格9（索引 {i}）")
            break
    
    if table_9 is None:
        print("ERROR: 未找到表格9")
        return
    
    # 获取当前表格信息
    rows = table_9.findall(f'{{{nsmap["w"]}}}tr')
    print(f"  当前行数: {len(rows)}")
    
    # 新增2022-2026年数据列
    new_columns = [
        # 表头
        ['2022', '2023', '2024', '2025', '2026'],
        # 美国政府债务/GDP
        ['120.6%', '120.4%', '121.7%', '123.3%', '127.8%'],
        # 美国居民杠杆率
        ['74.5%', '71.8%', '69.3%', '68.1%', '-'],
        # 中国政府债务/GDP
        ['77.3%', '84.1%', '90.4%', '99.2%', '106.9%'],
        # 中国企业杠杆率
        ['157.7%', '163.9%', '168.4%', '174.6%', '180.0%'],
        # 日本政府债务/GDP
        ['227.8%', '220.3%', '214.5%', '206.5%', '204.4%'],
    ]
    
    # 在表格末尾添加新列
    print("\n添加2022-2026年数据列...")
    for col_offset in range(5):
        col_idx = 5 + col_offset  # 从第6列开始（索引5）
        values = [row[col_offset] for row in new_columns]
        add_table_column(table_9, col_idx, values)
        print(f"  添加列: {values[0]}")
    
    # 保存
    print(f"\nSaving to {OUT_PATH}...")
    doc.save(OUT_PATH)
    
    # 验证
    verify_doc = Document(OUT_PATH)
    print(f"\nVerification:")
    print(f"  段落数: {len(verify_doc.paragraphs)}")
    print(f"  表格数: {len(verify_doc.tables)}")
    
    # 验证表格9
    for i, table in enumerate(verify_doc.tables):
        if i == 8:
            print(f"\n  表格9更新后:")
            for r, row in enumerate(table.rows):
                cells = [cell.text[:15] for cell in row.cells]
                print(f"    行{r}: {cells}")
    
    print(f"  ✅ Done!")


if __name__ == "__main__":
    main()
