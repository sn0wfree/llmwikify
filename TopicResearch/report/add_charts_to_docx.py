"""
直接向Word文档添加图表（修复图片未插入问题）
"""
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCX_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'
OUTPUT_PATH = '/home/ll/llmwikify/TopicResearch/report/复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.6.docx'
CHART_DIR = '/home/ll/llmwikify/TopicResearch/report/charts'

CHARTS = [
    ('chart_gdp_nominal.png', '图 1.1：名义 GDP 增速对比（2000-2026）', '2005-2007 年中国名义 GDP 增速从 15.6% 升至 23.2%，显著高于美国（6.7→4.8%）和日本（0.5→0.7%）。数据来源：ifind EDB。'),
    ('chart_gdp_real.png', '图 1.2：实际 GDP 增速对比（2000-2026）', '剔除价格因素后，中国实际 GDP 增速 11-14%，美国 2-3%，日本 1-2%。数据来源：ifind EDB。'),
    ('chart_cpi.png', '图 1.3：CPI 通胀对比（2000-2026）', '2005-2007 年中国 CPI 从 1.8% 升至 4.8%，美国维持 2-3%，日本长期通缩。数据来源：ifind EDB。'),
    ('chart_monetary.png', '图 1.4：中国货币政策背离', '2006-2007 年 PBOC 加息 6 次+提 RRR 10 次，但 M2 维持 17-19%。数据来源：ifind EDB。'),
    ('chart_term_spread.png', '图 1.5：市场利率与货币政策类型', '美国价格型、中国数量型→价格型、日本 YCC。数据来源：ifind EDB。'),
    ('chart_fx_trade.png', '图 1.6：汇率与贸易', '2005.7.21 汇改，DXY 从 120 跌至 80，中国顺差 1020→2639 亿$。数据来源：ifind EDB。'),
    ('chart_forex_reserve.png', '图 1.7：外汇储备/GDP', '中国外储/GDP 从 15% 升至 49%（2014），当前 17%。美国 0.3%，日本 20%。数据来源：ifind EDB。'),
    ('chart_impossible_trinity.png', '图 1.8：蒙代尔不可能三角', '固定汇率、资本自由流动、独立货币政策三者不可兼得。中国从 B 组合向 C 组合移动。'),
]

def main():
    print('Loading document...')
    doc = docx.Document(DOCX_PATH)
    print(f'Current paragraphs: {len(doc.paragraphs)}')
    
    # 在文档末尾添加图表章节
    print('\nAdding charts section...')
    
    # 添加分页符
    doc.add_page_break()
    
    # 添加章节标题
    p = doc.add_heading('附录 D：宏观经济背景图表', level=1)
    
    # 添加每个图表
    for filename, title, description in CHARTS:
        chart_path = os.path.join(CHART_DIR, filename)
        
        if not os.path.exists(chart_path):
            print(f'  ⚠️ Chart not found: {chart_path}')
            continue
        
        # 添加图表标题
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(10)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加图表
        p = doc.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(chart_path, width=Inches(6))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f'  ✅ Added: {title}')
        except Exception as e:
            print(f'  ❌ Failed to add {title}: {e}')
        
        # 添加说明文字
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.size = Pt(9)
    
    # 保存
    print(f'\nSaving to {OUTPUT_PATH}...')
    doc.save(OUTPUT_PATH)
    print(f'✅ Done! Saved to {OUTPUT_PATH}')
    
    # 验证
    doc2 = docx.Document(OUTPUT_PATH)
    rels = doc2.part.rels
    image_count = sum(1 for rel in rels.values() if 'image' in rel.reltype)
    print(f'\nVerification: {len(doc2.paragraphs)} paragraphs, {image_count} images')

if __name__ == '__main__':
    main()
