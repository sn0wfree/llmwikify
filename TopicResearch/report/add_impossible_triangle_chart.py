"""
添加不可能三角图表到 1.1.6 章节
"""
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.expanduser("~/llmwikify/TopicResearch/report")
CHARTS_DIR = os.path.join(BASE_DIR, "charts")
V29_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.9.docx")
OUT_PATH = os.path.join(BASE_DIR, "复盘2005–2007年有色金属超级周期与当前AI板块投资的结构性比较 v2.9.docx")


def main():
    print("Loading v2.9 document...")
    doc = Document(V29_PATH)
    print(f"  段落数: {len(doc.paragraphs)}")
    
    # 找到 1.1.6 章节位置
    for i, p in enumerate(doc.paragraphs):
        if "1.1.6 蒙代尔不可能三角" in p.text:
            print(f"  Found section 1.1.6 at paragraph {i}")
            
            # 找到表格后的位置（表格标题后）
            for j in range(i+1, min(i+10, len(doc.paragraphs))):
                if "表 1.1-3" in doc.paragraphs[j].text:
                    print(f"  Found table title at paragraph {j}")
                    
                    # 在表格后插入图表
                    # 找到表格结束位置
                    for k in range(j+1, min(j+5, len(doc.paragraphs))):
                        if "2005-2007年" in doc.paragraphs[k].text:
                            print(f"  Inserting chart after paragraph {k}")
                            
                            # 在 k+1 位置插入图表标题和图片
                            chart_path = os.path.join(CHARTS_DIR, "chart_impossible_trinity.png")
                            if os.path.exists(chart_path):
                                # 插入图表标题
                                new_para = doc.paragraphs[k]._element
                                from lxml import etree
                                nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                                
                                # 创建图表标题段落
                                title_p = etree.Element(f'{{{nsmap["w"]}}}p')
                                title_pPr = etree.SubElement(title_p, f'{{{nsmap["w"]}}}pPr')
                                title_jc = etree.SubElement(title_pPr, f'{{{nsmap["w"]}}}jc')
                                title_jc.set(f'{{{nsmap["w"]}}}val', 'center')
                                title_r = etree.SubElement(title_p, f'{{{nsmap["w"]}}}r')
                                title_rPr = etree.SubElement(title_r, f'{{{nsmap["w"]}}}rPr')
                                title_sz = etree.SubElement(title_rPr, f'{{{nsmap["w"]}}}sz')
                                title_sz.set(f'{{{nsmap["w"]}}}val', '18')
                                title_b = etree.SubElement(title_rPr, f'{{{nsmap["w"]}}}b')
                                title_t = etree.SubElement(title_r, f'{{{nsmap["w"]}}}t')
                                title_t.text = "图 1.3：蒙代尔不可能三角"
                                title_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                                
                                # 插入标题
                                new_para.addnext(title_p)
                                
                                # 创建图片段落
                                img_p = doc.add_paragraph()
                                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = img_p.add_run()
                                run.add_picture(chart_path, width=Inches(5.0))
                                
                                # 移动图片段落到正确位置
                                title_p.addnext(img_p._element)
                                
                                print(f"  ✅ Inserted chart")
                            break
                    break
            break
    
    # 保存
    print(f"\nSaving to {OUT_PATH}...")
    doc.save(OUT_PATH)
    
    # 验证
    verify_doc = Document(OUT_PATH)
    print(f"\nVerification:")
    print(f"  段落数: {len(verify_doc.paragraphs)}")
    print(f"  表格数: {len(verify_doc.tables)}")
    
    # 检查图片
    image_count = sum(1 for rel in verify_doc.part.rels.values() if 'image' in rel.reltype)
    print(f"  图片数: {image_count}")
    
    print(f"  ✅ Done!")


if __name__ == "__main__":
    main()
